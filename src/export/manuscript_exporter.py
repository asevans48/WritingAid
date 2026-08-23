"""Manuscript export functionality for various formats."""

import zipfile
from typing import List, Optional
from xml.sax.saxutils import escape as _xml_escape

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ebooklib import epub

from src.models.project import Chapter, Manuscript


# Office / document formats offered to the writer, in menu order.
# (key, label, extension). ``key`` is what ``export_document`` accepts.
DOCUMENT_FORMATS = [
    ("docx", "Word Document (.docx)", "docx"),
    ("rtf", "Rich Text (.rtf)", "rtf"),
    ("txt", "Plain Text (.txt)", "txt"),
    ("odt", "OpenDocument Text (.odt)", "odt"),
    ("ods", "OpenDocument Spreadsheet (.ods)", "ods"),
]


def _paragraphs(content: str) -> List[str]:
    """Split a chapter's plain-text content into paragraphs (blank-
    line separated), dropping empties. Falls back to line-splitting
    when the writer used single newlines."""
    text = (content or "").replace("\r\n", "\n").replace("\r", "\n")
    if not text.strip():
        return []
    blocks = [b.strip() for b in text.split("\n\n")]
    paras = [b for b in blocks if b]
    if len(paras) <= 1 and "\n" in text:
        # Single-newline prose — treat each non-empty line as a para.
        paras = [ln.strip() for ln in text.split("\n") if ln.strip()]
    return paras or [text.strip()]


class ManuscriptExporter:
    """Export manuscripts to various publishing formats."""

    def __init__(self, manuscript: Manuscript):
        """Initialize exporter with manuscript."""
        self.manuscript = manuscript

    # ------------------------------------------------------------------
    # Multi-format document export (Word / RTF / TXT / ODT / ODS)
    # ------------------------------------------------------------------
    def export_document(
        self,
        output_path: str,
        fmt: str,
        chapters: Optional[List[Chapter]] = None,
        title_page: bool = True,
    ) -> bool:
        """Export ``chapters`` (default: the whole manuscript) to
        ``output_path`` in ``fmt`` — one of ``docx``, ``rtf``,
        ``txt``, ``odt``, ``ods``. ``title_page`` adds the book title +
        author heading (used for whole-book exports, skipped for a
        single chapter). Returns True on success.

        Dependency-free: DOCX reuses python-docx (already a project
        dep); RTF, TXT, ODT and ODS are written directly, so no
        pandoc / LibreOffice / odfpy install is required."""
        chapters = (self.manuscript.chapters
                    if chapters is None else chapters)
        fmt = (fmt or "").lower()
        writers = {
            "docx": self._write_docx,
            "rtf": self._write_rtf,
            "txt": self._write_txt,
            "odt": self._write_odt,
            "ods": self._write_ods,
        }
        writer = writers.get(fmt)
        if writer is None:
            print(f"[export] unknown document format: {fmt}")
            return False
        try:
            return writer(output_path, chapters, title_page)
        except Exception as e:
            print(f"[export] {fmt} export failed: {e}")
            return False

    def _write_docx(
        self, output_path: str, chapters: List[Chapter],
        title_page: bool,
    ) -> bool:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        if title_page:
            title = doc.add_heading(self.manuscript.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if self.manuscript.author:
                author = doc.add_paragraph(
                    f"by {self.manuscript.author}")
                author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()
        for i, chapter in enumerate(chapters):
            doc.add_heading(chapter.title or f"Chapter {i + 1}", 1)
            for para in _paragraphs(chapter.content):
                doc.add_paragraph(para)
            if i < len(chapters) - 1:
                doc.add_page_break()
        doc.save(output_path)
        return True

    def _write_txt(
        self, output_path: str, chapters: List[Chapter],
        title_page: bool,
    ) -> bool:
        lines: List[str] = []
        if title_page:
            lines.append(self.manuscript.title or "Untitled")
            if self.manuscript.author:
                lines.append(f"by {self.manuscript.author}")
            lines.append("")
            lines.append("")
        for i, chapter in enumerate(chapters):
            lines.append(chapter.title or f"Chapter {i + 1}")
            lines.append("")
            for para in _paragraphs(chapter.content):
                lines.append(para)
                lines.append("")
            lines.append("")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines).rstrip() + "\n")
        return True

    def _write_rtf(
        self, output_path: str, chapters: List[Chapter],
        title_page: bool,
    ) -> bool:
        def esc(s: str) -> str:
            out = []
            for ch in (s or ""):
                o = ord(ch)
                if ch in "\\{}":
                    out.append("\\" + ch)
                elif o > 127:
                    # RTF signed 16-bit unicode escape + ASCII fallback.
                    out.append(
                        f"\\u{o if o < 32768 else o - 65536}?")
                else:
                    out.append(ch)
            return "".join(out)

        parts = [
            r"{\rtf1\ansi\ansicpg1252\deff0",
            r"{\fonttbl{\f0\froman Times New Roman;}}",
            r"\f0\fs24",
        ]
        if title_page:
            parts.append(
                r"\qc\b\fs36 " + esc(self.manuscript.title or "")
                + r"\b0\fs24\par")
            if self.manuscript.author:
                parts.append(
                    r"\qc by " + esc(self.manuscript.author) + r"\par")
            parts.append(r"\ql\par")
        for i, chapter in enumerate(chapters):
            parts.append(
                r"\b\fs30 "
                + esc(chapter.title or f"Chapter {i + 1}")
                + r"\b0\fs24\par")
            for para in _paragraphs(chapter.content):
                parts.append(esc(para) + r"\par")
            parts.append(r"\par")
        parts.append("}")
        with open(output_path, "w", encoding="ascii",
                   errors="ignore") as f:
            f.write("\n".join(parts))
        return True

    def _write_odt(
        self, output_path: str, chapters: List[Chapter],
        title_page: bool,
    ) -> bool:
        body: List[str] = []
        if title_page:
            body.append(
                '<text:h text:outline-level="1">'
                + _xml_escape(self.manuscript.title or "")
                + '</text:h>')
            if self.manuscript.author:
                body.append(
                    '<text:p>by '
                    + _xml_escape(self.manuscript.author)
                    + '</text:p>')
        for i, chapter in enumerate(chapters):
            body.append(
                '<text:h text:outline-level="1">'
                + _xml_escape(chapter.title or f"Chapter {i + 1}")
                + '</text:h>')
            for para in _paragraphs(chapter.content):
                body.append(
                    '<text:p>' + _xml_escape(para) + '</text:p>')
        content = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<office:document-content '
            'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:'
            'office:1.0" '
            'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:'
            'text:1.0" office:version="1.2">'
            '<office:body><office:text>'
            + "".join(body) +
            '</office:text></office:body></office:document-content>')
        self._write_odf_zip(
            output_path,
            "application/vnd.oasis.opendocument.text", content)
        return True

    def _write_ods(
        self, output_path: str, chapters: List[Chapter],
        title_page: bool,
    ) -> bool:
        def cell(value: str, numeric: bool = False) -> str:
            if numeric:
                return (
                    '<table:table-cell office:value-type="float" '
                    f'office:value="{value}"><text:p>{value}'
                    '</text:p></table:table-cell>')
            return (
                '<table:table-cell office:value-type="string">'
                '<text:p>' + _xml_escape(value)
                + '</text:p></table:table-cell>')

        def row(cells: List[str]) -> str:
            return ('<table:table-row>' + "".join(cells)
                    + '</table:table-row>')

        rows = [row([cell("Chapter"), cell("Words"),
                     cell("Content")])]
        for i, ch in enumerate(chapters):
            text = "\n\n".join(_paragraphs(ch.content))
            words = str(len((ch.content or "").split()))
            rows.append(row([
                cell(ch.title or f"Chapter {i + 1}"),
                cell(words, numeric=True),
                cell(text)]))
        content = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<office:document-content '
            'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:'
            'office:1.0" '
            'xmlns:table="urn:oasis:names:tc:opendocument:xmlns:'
            'table:1.0" '
            'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:'
            'text:1.0" office:version="1.2">'
            '<office:body><office:spreadsheet>'
            '<table:table table:name="Manuscript">'
            + "".join(rows) +
            '</table:table>'
            '</office:spreadsheet></office:body>'
            '</office:document-content>')
        self._write_odf_zip(
            output_path,
            "application/vnd.oasis.opendocument.spreadsheet",
            content)
        return True

    @staticmethod
    def _write_odf_zip(
        output_path: str, mimetype: str, content_xml: str,
    ) -> None:
        """Write a minimal but valid OpenDocument package: the
        ``mimetype`` entry MUST be first and stored uncompressed, then
        the manifest and content."""
        manifest = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<manifest:manifest '
            'xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:'
            'manifest:1.0" manifest:version="1.2">'
            f'<manifest:file-entry manifest:full-path="/" '
            f'manifest:media-type="{mimetype}"/>'
            '<manifest:file-entry manifest:full-path="content.xml" '
            'manifest:media-type="text/xml"/>'
            '</manifest:manifest>')
        with zipfile.ZipFile(
                output_path, "w", zipfile.ZIP_DEFLATED) as z:
            # mimetype first, STORED (uncompressed), no compression.
            zi = zipfile.ZipInfo("mimetype")
            zi.compress_type = zipfile.ZIP_STORED
            z.writestr(zi, mimetype)
            z.writestr("META-INF/manifest.xml", manifest)
            z.writestr("content.xml", content_xml)

    def export_to_docx(self, output_path: str) -> bool:
        """Export the whole manuscript as a Word document. Thin
        wrapper over the shared ``export_document`` writer so the
        existing callers (Export menu) keep working."""
        return self.export_document(output_path, "docx")

    def export_for_kindle(self, output_path: str) -> bool:
        """
        Export manuscript for Kindle Direct Publishing (EPUB format).

        KDP accepts EPUB, MOBI, and Word formats. EPUB is recommended.
        """
        try:
            book = epub.EpubBook()

            # Set metadata
            book.set_identifier(f'id-{self.manuscript.title}')
            book.set_title(self.manuscript.title)
            book.set_language('en')

            if self.manuscript.author:
                book.add_author(self.manuscript.author)

            # Create chapters
            epub_chapters = []

            for chapter in self.manuscript.chapters:
                # Create EPUB chapter
                c = epub.EpubHtml(
                    title=chapter.title,
                    file_name=f'chapter_{chapter.number}.xhtml',
                    lang='en'
                )

                # Format content
                content = f'<h1>{chapter.title}</h1>'
                # Convert plain text to HTML paragraphs
                paragraphs = chapter.content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        content += f'<p>{para.strip()}</p>'

                c.content = content

                book.add_item(c)
                epub_chapters.append(c)

            # Define Table of Contents
            book.toc = tuple(epub_chapters)

            # Add default NCX and Nav files
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            # Basic spine
            book.spine = ['nav'] + epub_chapters

            # Write EPUB file
            epub.write_epub(output_path, book)
            return True

        except Exception as e:
            print(f"Error exporting for Kindle: {e}")
            return False

    def export_for_barnes_noble(self, output_path: str) -> bool:
        """
        Export manuscript for Barnes & Noble Press (EPUB format).

        Barnes & Noble Press requires EPUB format.
        Similar to Kindle but with B&N-specific requirements.
        """
        # B&N Press uses standard EPUB format
        return self.export_for_kindle(output_path)

    def export_publisher_ready(self, output_path: str) -> bool:
        """
        Export manuscript in publisher-ready format (DOCX).

        Standard manuscript format:
        - 12pt Times New Roman
        - Double-spaced
        - 1-inch margins
        - Chapter titles on new pages
        - Page numbers
        """
        try:
            doc = Document()

            # Set up standard manuscript formatting
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Set paragraph formatting
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = 2.0  # Double-spaced
            paragraph_format.space_after = Pt(0)

            # Set margins (1 inch)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Title page
            title = doc.add_heading(self.manuscript.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if self.manuscript.author:
                author = doc.add_paragraph(f"by {self.manuscript.author}")
                author.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Word count
            word_count = doc.add_paragraph(f"Word Count: {self.manuscript.total_word_count:,}")
            word_count.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()

            # Add chapters
            for chapter in self.manuscript.chapters:
                # Chapter title - centered
                chapter_title = doc.add_heading(chapter.title, 1)
                chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Chapter content - double-spaced
                if chapter.content:
                    # Split into paragraphs
                    paragraphs = chapter.content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para.strip())
                            # Indent first line
                            p.paragraph_format.first_line_indent = Inches(0.5)

                doc.add_page_break()

            # Save document
            doc.save(output_path)
            return True

        except Exception as e:
            print(f"Error exporting publisher-ready manuscript: {e}")
            return False

    def export_outline_to_docx(self, output_path: str, include_notes: bool = True,
                               include_todos: bool = True) -> bool:
        """Export chapter plans/outlines as a Word document.

        This exports the chapter PLANS (not content) as a book outline.
        Useful for planning, sharing with editors, or keeping as reference.

        Args:
            output_path: Path to save the document
            include_notes: Whether to include chapter notes
            include_todos: Whether to include chapter todos

        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()

            # Set up document styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Title page
            title = doc.add_heading(f"{self.manuscript.title} - Book Outline", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if self.manuscript.author:
                author = doc.add_paragraph(f"by {self.manuscript.author}")
                author.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add outline summary
            doc.add_paragraph()
            summary = doc.add_paragraph()
            summary.add_run("Outline Summary").bold = True
            doc.add_paragraph(f"Total Chapters: {len(self.manuscript.chapters)}")

            # Count chapters with plans (check both legacy and new planning)
            chapters_with_plans = sum(
                1 for ch in self.manuscript.chapters
                if (ch.planning.outline and ch.planning.outline.strip()) or (ch.plan and ch.plan.strip())
            )
            doc.add_paragraph(f"Chapters with Plans: {chapters_with_plans}")

            # Count total todos
            total_todos = sum(len(ch.planning.todos) for ch in self.manuscript.chapters)
            completed_todos = sum(
                sum(1 for todo in ch.planning.todos if todo.completed)
                for ch in self.manuscript.chapters
            )
            if total_todos > 0:
                doc.add_paragraph(f"Writing Tasks: {completed_todos}/{total_todos} completed")

            doc.add_page_break()

            # Table of Contents style overview
            toc_heading = doc.add_heading("Chapter Overview", 1)

            for i, chapter in enumerate(self.manuscript.chapters, 1):
                # Check both planning.outline and legacy plan field
                outline = chapter.planning.outline or chapter.plan
                has_plan = "✓" if (outline and outline.strip()) else "○"
                word_count = chapter.word_count or 0
                status = f"({word_count} words)" if word_count > 0 else "(not started)"
                doc.add_paragraph(f"{has_plan} Chapter {i}: {chapter.title} {status}")

            doc.add_page_break()

            # Detailed chapter plans
            doc.add_heading("Detailed Chapter Plans", 1)

            for chapter in self.manuscript.chapters:
                # Chapter heading
                chapter_heading = doc.add_heading(f"Chapter {chapter.number}: {chapter.title}", 2)

                # Chapter metadata
                meta = doc.add_paragraph()
                meta.add_run(f"Word Count: ").bold = True
                meta.add_run(f"{chapter.word_count or 0} words")

                # Description (new field)
                if chapter.planning.description and chapter.planning.description.strip():
                    doc.add_paragraph()
                    desc_label = doc.add_paragraph()
                    desc_label.add_run("Description:").bold = True

                    desc_paragraphs = chapter.planning.description.strip().split('\n')
                    for para in desc_paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para.strip())
                            p.paragraph_format.left_indent = Inches(0.25)

                # Chapter outline/plan (check planning.outline first, then legacy plan)
                outline = chapter.planning.outline or chapter.plan
                doc.add_paragraph()
                plan_label = doc.add_paragraph()
                plan_label.add_run("Chapter Outline:").bold = True

                if outline and outline.strip():
                    # Split plan into paragraphs for better formatting
                    plan_paragraphs = outline.strip().split('\n')
                    for para in plan_paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para.strip())
                            p.paragraph_format.left_indent = Inches(0.25)
                else:
                    no_plan = doc.add_paragraph("(No outline written yet)")
                    no_plan.italic = True
                    no_plan.paragraph_format.left_indent = Inches(0.25)

                # Todos (new field)
                if include_todos and chapter.planning.todos:
                    doc.add_paragraph()
                    todos_label = doc.add_paragraph()
                    todos_label.add_run("Writing Tasks:").bold = True

                    for todo in chapter.planning.todos:
                        check = "☑" if todo.completed else "☐"
                        priority_marker = ""
                        if todo.priority == "high":
                            priority_marker = " [HIGH]"
                        elif todo.priority == "low":
                            priority_marker = " [low]"

                        p = doc.add_paragraph(f"{check} {todo.text}{priority_marker}")
                        p.paragraph_format.left_indent = Inches(0.25)

                # Chapter notes (check planning.notes first, then legacy notes)
                notes = chapter.planning.notes_as_text or chapter.notes
                if include_notes and notes and notes.strip():
                    doc.add_paragraph()
                    notes_label = doc.add_paragraph()
                    notes_label.add_run("Notes:").bold = True

                    notes_paragraphs = notes.strip().split('\n')
                    for para in notes_paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para.strip())
                            p.paragraph_format.left_indent = Inches(0.25)

                # Add separator between chapters
                doc.add_paragraph()
                doc.add_paragraph("─" * 50)
                doc.add_paragraph()

            # Save document
            doc.save(output_path)
            return True

        except Exception as e:
            print(f"Error exporting outline to DOCX: {e}")
            return False

    def get_manuscript_statistics(self) -> dict:
        """Get manuscript statistics for submission purposes."""
        total_words = sum(chapter.word_count for chapter in self.manuscript.chapters)
        total_chapters = len(self.manuscript.chapters)

        # Estimate pages (250 words per page is standard)
        estimated_pages = total_words / 250

        return {
            'total_words': total_words,
            'total_chapters': total_chapters,
            'estimated_pages': int(estimated_pages),
            'title': self.manuscript.title,
            'author': self.manuscript.author
        }
