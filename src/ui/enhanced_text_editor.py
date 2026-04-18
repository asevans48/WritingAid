"""Enhanced text editor with context menu, Word support, headers, spell checking, and TTS."""

from PyQt6.QtWidgets import (
    QTextEdit, QMenu, QDialog, QVBoxLayout, QLabel,
    QListWidget, QPushButton, QTextBrowser, QHBoxLayout,
    QMessageBox, QComboBox, QToolTip, QWidgetAction, QSlider,
    QGroupBox
)
from PyQt6.QtCore import pyqtSignal, Qt, QTimer, QPoint, QEvent
from PyQt6.QtGui import (
    QAction, QTextCursor, QTextCharFormat, QColor, QFont,
    QTextBlockFormat, QSyntaxHighlighter, QTextDocument,
    QTextBlockUserData, QMouseEvent
)
from typing import Optional, Callable, List, Set, Dict, Tuple
from dataclasses import dataclass
from enum import Enum
from collections import Counter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

# Markdown utilities
from src.utils.markdown_editor import (
    MarkdownStyle, STYLE_NAMES, STYLE_TO_NAME, get_line_style,
    apply_heading_to_line, strip_markdown
)

# TTS imports (lazy loaded)
try:
    from src.services.tts_service import get_tts_service, TTSEngine
    from src.services.tts_document_generator import (
        TTSDocumentGenerator, TTSDocumentConfig, TTSFormat,
        SpeakerConfig, get_tts_output_dir
    )
    TTS_AVAILABLE = True
except ImportError:
    TTS_AVAILABLE = False


class HeadingLevel(Enum):
    """Heading levels for document structure."""
    NORMAL = 0       # Regular paragraph
    TITLE = 1        # Document title (Heading 0 in Word)
    HEADING_1 = 2    # Chapter/Major section
    HEADING_2 = 3    # Section
    HEADING_3 = 4    # Subsection
    HEADING_4 = 5    # Minor subsection


class BlockData(QTextBlockUserData):
    """Custom block data to store heading level and other metadata."""

    def __init__(self, heading_level: HeadingLevel = HeadingLevel.NORMAL):
        super().__init__()
        self.heading_level = heading_level


class CheckMode(Enum):
    """Modes for writing checks."""
    ON_DEMAND = "on_demand"  # Only runs on recheck button click (green)
    REALTIME = "realtime"    # Runs automatically as you type (blue)
    OFF = "off"              # Disabled (white/default)


class ErrorType(Enum):
    """Types of writing errors."""
    SPELLING = "spelling"
    GRAMMAR = "grammar"
    OVERUSE = "overuse"
    OVERUSED_PHRASE = "overused_phrase"


@dataclass
class WritingError:
    """Represents a writing error with details."""
    error_type: ErrorType
    word: str
    start: int
    end: int
    reason: str
    suggestions: List[str]


class SpellChecker:
    """Spell checker using nltk words corpus or pyenchant."""

    def __init__(self):
        """Initialize spell checker."""
        self._dictionary: Set[str] = set()
        self._custom_words: Set[str] = set()
        self._use_enchant = False
        self._enchant_dict = None
        self._use_pyspellchecker = False
        self._pyspellchecker = None
        self._misspelling_reasons: Dict[str, str] = {}
        self._load_dictionary()

    def _load_dictionary(self):
        """Load dictionary - try multiple sources."""
        # First try pyspellchecker (pure Python, works everywhere)
        try:
            from spellchecker import SpellChecker as PySpellChecker
            self._pyspellchecker = PySpellChecker()
            self._use_pyspellchecker = True
            print("Spell checker: Using pyspellchecker")
            return
        except ImportError:
            self._use_pyspellchecker = False
            self._pyspellchecker = None

        # Try pyenchant (requires system library)
        try:
            import enchant
            # Check if en_US dictionary is available
            if enchant.dict_exists("en_US"):
                self._enchant_dict = enchant.Dict("en_US")
                self._use_enchant = True
                print("Spell checker: Using pyenchant")
                return
            elif enchant.dict_exists("en"):
                self._enchant_dict = enchant.Dict("en")
                self._use_enchant = True
                print("Spell checker: Using pyenchant (en)")
                return
            else:
                print("Spell checker: pyenchant installed but no English dictionary found")
        except ImportError:
            pass
        except Exception as e:
            print(f"Spell checker: pyenchant error: {e}")

        # Try nltk words corpus
        try:
            import nltk
            try:
                from nltk.corpus import words
                self._dictionary = set(w.lower() for w in words.words())
                print(f"Spell checker: Using nltk ({len(self._dictionary)} words)")
                return
            except LookupError:
                # Download if not present
                nltk.download('words', quiet=True)
                from nltk.corpus import words
                self._dictionary = set(w.lower() for w in words.words())
                print(f"Spell checker: Using nltk ({len(self._dictionary)} words)")
                return
        except ImportError:
            pass

        # Fallback: comprehensive built-in dictionary
        print("Spell checker: Using built-in dictionary (install pyenchant or nltk for better results)")
        self._dictionary = self._get_comprehensive_dictionary()

    def _get_comprehensive_dictionary(self) -> Set[str]:
        """Get a comprehensive built-in dictionary of common English words."""
        # Top 10,000+ most common English words covering most writing needs
        # This is a subset - the full list would be much larger
        common_words = {
            # Articles, prepositions, conjunctions
            "a", "an", "the", "and", "but", "or", "nor", "for", "yet", "so",
            "at", "by", "from", "in", "into", "of", "off", "on", "onto", "out",
            "over", "to", "up", "with", "as", "if", "than", "that", "though",
            "after", "before", "since", "until", "while", "about", "above",
            "across", "against", "along", "among", "around", "behind", "below",
            "beneath", "beside", "between", "beyond", "down", "during", "except",
            "inside", "near", "outside", "through", "toward", "towards", "under",
            "upon", "within", "without",

            # Pronouns
            "i", "me", "my", "mine", "myself", "you", "your", "yours", "yourself",
            "yourselves", "he", "him", "his", "himself", "she", "her", "hers",
            "herself", "it", "its", "itself", "we", "us", "our", "ours", "ourselves",
            "they", "them", "their", "theirs", "themselves", "who", "whom", "whose",
            "which", "what", "this", "that", "these", "those", "each", "every",
            "either", "neither", "both", "all", "any", "some", "none", "one",
            "anyone", "everyone", "someone", "no one", "nobody", "anybody",
            "everybody", "somebody", "anything", "everything", "something", "nothing",

            # Common verbs (base forms and conjugations)
            "be", "am", "is", "are", "was", "were", "been", "being",
            "have", "has", "had", "having",
            "do", "does", "did", "doing", "done",
            "say", "says", "said", "saying",
            "get", "gets", "got", "gotten", "getting",
            "make", "makes", "made", "making",
            "go", "goes", "went", "gone", "going",
            "know", "knows", "knew", "known", "knowing",
            "take", "takes", "took", "taken", "taking",
            "see", "sees", "saw", "seen", "seeing",
            "come", "comes", "came", "coming",
            "think", "thinks", "thought", "thinking",
            "look", "looks", "looked", "looking",
            "want", "wants", "wanted", "wanting",
            "give", "gives", "gave", "given", "giving",
            "use", "uses", "used", "using",
            "find", "finds", "found", "finding",
            "tell", "tells", "told", "telling",
            "ask", "asks", "asked", "asking",
            "work", "works", "worked", "working",
            "seem", "seems", "seemed", "seeming",
            "feel", "feels", "felt", "feeling",
            "try", "tries", "tried", "trying",
            "leave", "leaves", "left", "leaving",
            "call", "calls", "called", "calling",
            "keep", "keeps", "kept", "keeping",
            "let", "lets", "letting",
            "begin", "begins", "began", "begun", "beginning",
            "show", "shows", "showed", "shown", "showing",
            "hear", "hears", "heard", "hearing",
            "play", "plays", "played", "playing",
            "run", "runs", "ran", "running",
            "move", "moves", "moved", "moving",
            "live", "lives", "lived", "living",
            "believe", "believes", "believed", "believing",
            "hold", "holds", "held", "holding",
            "bring", "brings", "brought", "bringing",
            "happen", "happens", "happened", "happening",
            "write", "writes", "wrote", "written", "writing",
            "provide", "provides", "provided", "providing",
            "sit", "sits", "sat", "sitting",
            "stand", "stands", "stood", "standing",
            "lose", "loses", "lost", "losing",
            "pay", "pays", "paid", "paying",
            "meet", "meets", "met", "meeting",
            "include", "includes", "included", "including",
            "continue", "continues", "continued", "continuing",
            "set", "sets", "setting",
            "learn", "learns", "learned", "learnt", "learning",
            "change", "changes", "changed", "changing",
            "lead", "leads", "led", "leading",
            "understand", "understands", "understood", "understanding",
            "watch", "watches", "watched", "watching",
            "follow", "follows", "followed", "following",
            "stop", "stops", "stopped", "stopping",
            "create", "creates", "created", "creating",
            "speak", "speaks", "spoke", "spoken", "speaking",
            "read", "reads", "reading",
            "spend", "spends", "spent", "spending",
            "grow", "grows", "grew", "grown", "growing",
            "open", "opens", "opened", "opening",
            "walk", "walks", "walked", "walking",
            "win", "wins", "won", "winning",
            "offer", "offers", "offered", "offering",
            "remember", "remembers", "remembered", "remembering",
            "love", "loves", "loved", "loving",
            "consider", "considers", "considered", "considering",
            "appear", "appears", "appeared", "appearing",
            "buy", "buys", "bought", "buying",
            "wait", "waits", "waited", "waiting",
            "serve", "serves", "served", "serving",
            "die", "dies", "died", "dying",
            "send", "sends", "sent", "sending",
            "expect", "expects", "expected", "expecting",
            "build", "builds", "built", "building",
            "stay", "stays", "stayed", "staying",
            "fall", "falls", "fell", "fallen", "falling",
            "cut", "cuts", "cutting",
            "reach", "reaches", "reached", "reaching",
            "kill", "kills", "killed", "killing",
            "remain", "remains", "remained", "remaining",
            "suggest", "suggests", "suggested", "suggesting",
            "raise", "raises", "raised", "raising",
            "pass", "passes", "passed", "passing",
            "sell", "sells", "sold", "selling",
            "require", "requires", "required", "requiring",
            "report", "reports", "reported", "reporting",
            "decide", "decides", "decided", "deciding",
            "pull", "pulls", "pulled", "pulling",
            "develop", "develops", "developed", "developing",
            "wish", "wishes", "wished", "wishing",
            "drop", "drops", "dropped", "dropping",
            "push", "pushes", "pushed", "pushing",
            "wear", "wears", "wore", "worn", "wearing",
            "cause", "causes", "caused", "causing",
            "return", "returns", "returned", "returning",
            "realize", "realizes", "realized", "realizing",
            "explain", "explains", "explained", "explaining",
            "hope", "hopes", "hoped", "hoping",
            "drive", "drives", "drove", "driven", "driving",
            "carry", "carries", "carried", "carrying",
            "allow", "allows", "allowed", "allowing",
            "suppose", "supposes", "supposed", "supposing",
            "agree", "agrees", "agreed", "agreeing",
            "help", "helps", "helped", "helping",
            "put", "puts", "putting",
            "mean", "means", "meant", "meaning",
            "might", "must", "can", "could", "would", "should", "shall", "will",

            # Common nouns
            "time", "year", "people", "way", "day", "man", "woman", "child",
            "children", "world", "life", "hand", "part", "place", "case", "week",
            "company", "system", "program", "question", "work", "government",
            "number", "night", "point", "home", "water", "room", "mother",
            "area", "money", "story", "fact", "month", "lot", "right", "study",
            "book", "eye", "job", "word", "business", "issue", "side", "kind",
            "head", "house", "service", "friend", "father", "power", "hour",
            "game", "line", "end", "member", "law", "car", "city", "community",
            "name", "president", "team", "minute", "idea", "kid", "body",
            "information", "back", "parent", "face", "others", "level", "office",
            "door", "health", "person", "art", "war", "history", "party",
            "result", "change", "morning", "reason", "research", "girl", "guy",
            "moment", "air", "teacher", "force", "education", "foot", "feet",
            "boy", "age", "policy", "process", "music", "market", "sense",
            "nation", "plan", "college", "interest", "death", "experience",
            "effect", "use", "class", "control", "care", "field", "development",
            "role", "effort", "rate", "heart", "drug", "show", "leader", "light",
            "voice", "wife", "police", "mind", "difference", "period", "value",
            "building", "action", "authority", "model", "husband", "support",
            "event", "picture", "evidence", "product", "truth", "position",
            "reason", "technology", "attention", "rule", "concern", "news",
            "player", "price", "chance", "view", "practice", "past", "doctor",
            "wall", "patient", "nature", "section", "fire", "society", "hair",
            "form", "activity", "country", "road", "table", "court", "form",
            "record", "situation", "food", "student", "material", "paper",
            "group", "defense", "manager", "character", "analysis", "access",
            "season", "today", "film", "growth", "income", "statement", "peace",
            "century", "space", "bed", "order", "presence", "opportunity",
            "performance", "relationship", "author", "response", "language",
            "love", "knowledge", "friend", "size", "bank", "condition", "base",
            "thought", "article", "look", "pressure", "behavior", "movement",
            "century", "sun", "blood", "series", "land", "project", "choice",
            "oil", "step", "amount", "church", "town", "attention", "theory",
            "director", "figure", "attack", "design", "loss", "purpose",
            "image", "standard", "approach", "agreement", "decision", "direction",

            # Common adjectives
            "good", "new", "first", "last", "long", "great", "little", "own",
            "other", "old", "right", "big", "high", "different", "small", "large",
            "next", "early", "young", "important", "few", "public", "bad", "same",
            "able", "human", "local", "late", "hard", "major", "better", "economic",
            "strong", "possible", "whole", "free", "military", "true", "federal",
            "international", "full", "special", "easy", "clear", "recent",
            "certain", "personal", "open", "red", "difficult", "available",
            "likely", "short", "single", "medical", "current", "wrong", "private",
            "past", "foreign", "fine", "common", "poor", "natural", "significant",
            "similar", "hot", "dead", "central", "happy", "serious", "ready",
            "simple", "left", "physical", "general", "environmental", "financial",
            "blue", "democratic", "dark", "various", "entire", "close", "legal",
            "religious", "cold", "final", "main", "green", "nice", "huge", "popular",
            "traditional", "cultural", "beautiful", "real", "safe", "light",
            "sure", "white", "black", "deep", "wide", "complete", "wild",

            # Common adverbs
            "not", "also", "very", "often", "however", "too", "usually", "really",
            "early", "never", "always", "sometimes", "together", "likely", "simply",
            "generally", "instead", "actually", "remember", "likely", "perhaps",
            "quickly", "slowly", "finally", "almost", "above", "hard", "especially",
            "probably", "already", "below", "directly", "therefore", "once",
            "ever", "else", "rather", "nearly", "certainly", "clearly", "soon",
            "indeed", "easily", "eventually", "exactly", "highly", "suddenly",

            # Writing-related words
            "chapter", "story", "plot", "scene", "dialogue", "narrative",
            "protagonist", "antagonist", "setting", "conflict", "resolution",
            "climax", "theme", "motif", "symbol", "foreshadowing", "flashback",
            "narrator", "manuscript", "draft", "edit", "revision", "paragraph",
            "sentence", "phrase", "metaphor", "simile", "imagery", "tone",
            "voice", "style", "genre", "fiction", "nonfiction", "novel",
            "novella", "screenplay", "script", "outline", "arc", "pacing",
            "tension", "suspense", "mystery", "fantasy", "romance", "thriller",
            "horror", "drama", "comedy", "tragedy", "epic", "saga", "trilogy",
            "sequel", "prequel", "spinoff", "backstory", "worldbuilding",
            "characterization", "motivation", "emotion", "feeling", "mood",

            # Numbers as words
            "zero", "one", "two", "three", "four", "five", "six", "seven",
            "eight", "nine", "ten", "eleven", "twelve", "thirteen", "fourteen",
            "fifteen", "sixteen", "seventeen", "eighteen", "nineteen", "twenty",
            "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety",
            "hundred", "thousand", "million", "billion",

            # Time-related
            "today", "tomorrow", "yesterday", "now", "then", "soon", "later",
            "morning", "afternoon", "evening", "night", "midnight", "noon",
            "second", "minute", "hour", "day", "week", "month", "year",
            "decade", "century", "monday", "tuesday", "wednesday", "thursday",
            "friday", "saturday", "sunday", "january", "february", "march",
            "april", "may", "june", "july", "august", "september", "october",
            "november", "december", "spring", "summer", "autumn", "fall", "winter",

            # Contractions (without apostrophe for matching)
            "dont", "wont", "cant", "shouldnt", "wouldnt", "couldnt", "didnt",
            "doesnt", "isnt", "arent", "wasnt", "werent", "hasnt", "havent",
            "hadnt", "im", "youre", "hes", "shes", "its", "were", "theyre",
            "ive", "youve", "weve", "theyve", "id", "youd", "hed", "shed",
            "wed", "theyd", "ill", "youll", "hell", "shell", "well", "theyll",
            "thats", "whats", "whos", "heres", "theres", "wheres", "lets",

            # Additional common words
            "yes", "no", "ok", "okay", "please", "thank", "thanks", "sorry",
            "hello", "goodbye", "hey", "oh", "well", "just", "like", "really",
            "thing", "things", "stuff", "something", "anything", "nothing",
            "everything", "someone", "anyone", "nobody", "everybody", "somewhere",
        }
        return common_words

    def check_word(self, word: str) -> Tuple[bool, str]:
        """Check if a word is spelled correctly.

        Returns:
            Tuple of (is_correct, reason_if_incorrect)
        """
        # Skip very short words
        if len(word) <= 1:
            return True, ""

        # Skip words that are all caps (acronyms)
        if word.isupper() and len(word) > 1:
            return True, ""

        # Skip words starting with capital (proper nouns)
        if word[0].isupper():
            return True, ""

        # Skip numbers and words with numbers
        if any(c.isdigit() for c in word):
            return True, ""

        # Skip words with special characters
        if not word.isalpha():
            return True, ""

        word_lower = word.lower()

        # Check custom words first
        if word_lower in self._custom_words:
            return True, ""

        # Check with pyspellchecker if available (preferred)
        if self._use_pyspellchecker and self._pyspellchecker:
            if self._pyspellchecker.known([word_lower]):
                return True, ""
            else:
                suggestions = list(self._pyspellchecker.candidates(word_lower) or [])[:3]
                if suggestions:
                    return False, f"Possible misspelling. Did you mean: {', '.join(suggestions)}?"
                return False, "Word not found in dictionary"

        # Check with enchant if available
        if self._use_enchant and self._enchant_dict:
            if self._enchant_dict.check(word):
                return True, ""
            else:
                suggestions = self._enchant_dict.suggest(word)[:3]
                if suggestions:
                    return False, f"Possible misspelling. Did you mean: {', '.join(suggestions)}?"
                return False, "Word not found in dictionary"

        # Check built-in dictionary
        if word_lower in self._dictionary:
            return True, ""

        # Check common suffixes
        if self._check_with_suffixes(word_lower):
            return True, ""

        return False, "Word not found in dictionary"

    def _check_with_suffixes(self, word: str) -> bool:
        """Check if word with common suffixes is valid."""
        # Common suffixes to try removing
        suffixes = ['s', 'es', 'ed', 'ing', 'ly', 'er', 'est', 'ness', 'ment',
                   'tion', 'sion', 'able', 'ible', 'ful', 'less', 'ous', 'ive']

        for suffix in suffixes:
            if word.endswith(suffix):
                root = word[:-len(suffix)]
                if len(root) >= 2 and root in self._dictionary:
                    return True
                # Handle doubling (e.g., running -> run)
                if len(root) >= 3 and root[-1] == root[-2] and root[:-1] in self._dictionary:
                    return True

        # Check -ing with e-dropping (e.g., making -> make)
        if word.endswith('ing') and len(word) > 4:
            root = word[:-3] + 'e'
            if root in self._dictionary:
                return True

        # Check -ed with e-dropping (e.g., loved -> love)
        if word.endswith('ed') and len(word) > 3:
            root = word[:-1]  # e.g., loved -> love
            if root in self._dictionary:
                return True
            root = word[:-2]  # e.g., walked -> walk
            if root in self._dictionary:
                return True

        return False

    def check_word_simple(self, word: str) -> bool:
        """Simple check returning just True/False."""
        is_correct, _ = self.check_word(word)
        return is_correct

    def get_reason(self, word: str) -> str:
        """Get the reason why a word is flagged."""
        _, reason = self.check_word(word)
        return reason

    def suggest(self, word: str) -> List[str]:
        """Get spelling suggestions for a word."""
        if self._use_pyspellchecker and self._pyspellchecker:
            candidates = self._pyspellchecker.candidates(word.lower())
            return list(candidates)[:5] if candidates else []

        if self._use_enchant and self._enchant_dict:
            return self._enchant_dict.suggest(word)[:5]

        # Simple suggestions based on edit distance
        suggestions = []
        word_lower = word.lower()

        for dict_word in self._dictionary:
            if abs(len(dict_word) - len(word_lower)) <= 2:
                distance = self._edit_distance(word_lower, dict_word)
                if distance <= 2:
                    suggestions.append((distance, dict_word))

        suggestions.sort(key=lambda x: x[0])
        return [w for _, w in suggestions[:5]]

    def _edit_distance(self, s1: str, s2: str) -> int:
        """Calculate Levenshtein edit distance."""
        if len(s1) < len(s2):
            return self._edit_distance(s2, s1)

        if len(s2) == 0:
            return len(s1)

        prev_row = range(len(s2) + 1)
        for i, c1 in enumerate(s1):
            curr_row = [i + 1]
            for j, c2 in enumerate(s2):
                insertions = prev_row[j + 1] + 1
                deletions = curr_row[j] + 1
                substitutions = prev_row[j] + (c1 != c2)
                curr_row.append(min(insertions, deletions, substitutions))
            prev_row = curr_row

        return prev_row[-1]

    def add_word(self, word: str):
        """Add a word to the custom dictionary."""
        self._custom_words.add(word.lower())

    def is_available(self) -> bool:
        """Check if spell checking is properly available."""
        return self._use_pyspellchecker or self._use_enchant or len(self._dictionary) > 100


class OveruseDetector:
    """Detects overused words and phrases, providing synonym suggestions."""

    # Words to ignore (common words that are expected to repeat)
    IGNORE_WORDS = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'been',
        'be', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
        'could', 'should', 'may', 'might', 'must', 'shall', 'can', 'need',
        'it', 'its', 'this', 'that', 'these', 'those', 'i', 'you', 'he',
        'she', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your',
        'his', 'her', 'our', 'their', 'what', 'which', 'who', 'whom', 'whose',
        'if', 'then', 'else', 'when', 'where', 'why', 'how', 'all', 'each',
        'every', 'both', 'few', 'more', 'most', 'other', 'some', 'such', 'no',
        'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 'just',
        'also', 'now', 'here', 'there', 'still', 'already', 'even', 'again',
    }

    # Common phrases to ignore (expected to repeat)
    IGNORE_PHRASES = {
        'he said', 'she said', 'i said', 'they said',
        'it was', 'there was', 'there were', 'it is',
        'i was', 'he was', 'she was', 'they were',
        'i am', 'he is', 'she is', 'they are',
        'in the', 'on the', 'at the', 'to the', 'of the', 'for the',
        'and the', 'but the', 'with the', 'from the',
        'i had', 'he had', 'she had', 'they had',
        'i have', 'he has', 'she has', 'they have',
        'as if', 'as though', 'even if', 'even though',
        'out of', 'in front of', 'on top of', 'because of',
        'a lot', 'a lot of', 'some of', 'all of', 'one of',
    }

    # Common synonyms for overused words
    SYNONYMS = {
        'said': ['replied', 'responded', 'answered', 'stated', 'remarked', 'declared', 'exclaimed', 'muttered', 'whispered', 'shouted'],
        'walked': ['strolled', 'ambled', 'strode', 'marched', 'trudged', 'sauntered', 'paced', 'wandered', 'meandered', 'trekked'],
        'looked': ['glanced', 'gazed', 'stared', 'peered', 'observed', 'watched', 'examined', 'studied', 'surveyed', 'inspected'],
        'went': ['traveled', 'journeyed', 'proceeded', 'headed', 'moved', 'advanced', 'continued', 'departed', 'left', 'ventured'],
        'got': ['obtained', 'acquired', 'received', 'gained', 'secured', 'fetched', 'retrieved', 'collected', 'gathered', 'earned'],
        'good': ['excellent', 'great', 'fine', 'wonderful', 'superb', 'outstanding', 'remarkable', 'exceptional', 'splendid', 'admirable'],
        'bad': ['terrible', 'awful', 'dreadful', 'poor', 'inferior', 'substandard', 'unpleasant', 'disagreeable', 'unfavorable', 'negative'],
        'big': ['large', 'huge', 'enormous', 'massive', 'immense', 'vast', 'substantial', 'considerable', 'significant', 'grand'],
        'small': ['tiny', 'little', 'minute', 'miniature', 'compact', 'petite', 'modest', 'slight', 'minor', 'diminutive'],
        'nice': ['pleasant', 'agreeable', 'delightful', 'lovely', 'enjoyable', 'charming', 'appealing', 'attractive', 'pleasing', 'wonderful'],
        'happy': ['joyful', 'elated', 'delighted', 'pleased', 'content', 'cheerful', 'thrilled', 'ecstatic', 'overjoyed', 'blissful'],
        'sad': ['unhappy', 'sorrowful', 'melancholy', 'dejected', 'downcast', 'gloomy', 'miserable', 'despondent', 'heartbroken', 'forlorn'],
        'very': ['extremely', 'highly', 'incredibly', 'remarkably', 'exceptionally', 'particularly', 'especially', 'tremendously', 'immensely', 'greatly'],
        'really': ['truly', 'genuinely', 'actually', 'indeed', 'certainly', 'definitely', 'absolutely', 'positively', 'undoubtedly', 'unquestionably'],
        'thing': ['object', 'item', 'element', 'aspect', 'matter', 'subject', 'entity', 'article', 'component', 'detail'],
        'things': ['objects', 'items', 'elements', 'aspects', 'matters', 'subjects', 'entities', 'articles', 'components', 'details'],
        'stuff': ['material', 'items', 'belongings', 'possessions', 'goods', 'objects', 'articles', 'equipment', 'supplies', 'gear'],
        'like': ['such as', 'similar to', 'resembling', 'comparable to', 'akin to'],
        'beautiful': ['gorgeous', 'stunning', 'lovely', 'attractive', 'exquisite', 'magnificent', 'elegant', 'radiant', 'striking', 'breathtaking'],
        'fast': ['quick', 'rapid', 'swift', 'speedy', 'hasty', 'brisk', 'fleet', 'expeditious', 'prompt', 'nimble'],
        'slow': ['sluggish', 'leisurely', 'unhurried', 'gradual', 'deliberate', 'measured', 'plodding', 'languid', 'lethargic', 'torpid'],
        'think': ['believe', 'consider', 'suppose', 'reckon', 'assume', 'imagine', 'presume', 'surmise', 'speculate', 'ponder'],
        'know': ['understand', 'realize', 'recognize', 'comprehend', 'perceive', 'grasp', 'fathom', 'discern', 'appreciate', 'acknowledge'],
        'see': ['observe', 'notice', 'perceive', 'witness', 'spot', 'detect', 'discern', 'view', 'behold', 'glimpse'],
        'make': ['create', 'produce', 'construct', 'build', 'form', 'generate', 'manufacture', 'craft', 'fabricate', 'fashion'],
        'take': ['grab', 'seize', 'grasp', 'snatch', 'acquire', 'obtain', 'secure', 'capture', 'collect', 'gather'],
        'come': ['arrive', 'approach', 'appear', 'emerge', 'materialize', 'surface', 'show up', 'turn up', 'reach', 'enter'],
        'give': ['provide', 'offer', 'present', 'supply', 'grant', 'bestow', 'donate', 'contribute', 'deliver', 'hand over'],
        'tell': ['inform', 'notify', 'advise', 'explain', 'describe', 'relate', 'recount', 'narrate', 'disclose', 'reveal'],
        'feel': ['sense', 'experience', 'perceive', 'detect', 'notice', 'undergo', 'encounter', 'endure', 'suffer', 'enjoy'],
        'eyes': ['gaze', 'vision', 'sight', 'glance', 'stare', 'look'],
        'face': ['expression', 'countenance', 'features', 'visage', 'appearance', 'complexion'],
        'began': ['started', 'commenced', 'initiated', 'launched', 'embarked', 'undertook', 'opened', 'set out'],
        'suddenly': ['abruptly', 'unexpectedly', 'instantly', 'immediately', 'swiftly', 'rapidly', 'quickly', 'all at once'],
    }

    def __init__(self, threshold: int = 3, phrase_threshold: int = 2, window_size: int = 500):
        """Initialize overuse detector.

        Args:
            threshold: Number of times a word must appear to be flagged
            phrase_threshold: Number of times a phrase must appear to be flagged
            window_size: Number of words to consider for frequency analysis
        """
        self.threshold = threshold
        self.phrase_threshold = phrase_threshold
        self.window_size = window_size
        self._ignored_words: Set[str] = set()
        self._ignored_phrases: Set[str] = set()
        # Cache for word counts to avoid recalculating for every word
        self._cached_counts: Dict[str, int] = {}
        self._cached_text_hash: int = 0
        # Cache for phrase counts (2, 3, 4, 5-word ngrams)
        self._cached_phrase_counts: Dict[str, int] = {}
        # Cache phrase positions for highlighting: phrase -> [(start, end), ...]
        self._cached_phrase_positions: Dict[str, List[Tuple[int, int]]] = {}

    def _extract_ngrams(self, words: List[str], n: int) -> List[str]:
        """Extract n-grams from a list of words."""
        if len(words) < n:
            return []
        return [' '.join(words[i:i+n]) for i in range(len(words) - n + 1)]

    def update_cache(self, text: str):
        """Update the cached word and phrase counts for a text."""
        text_hash = hash(text)
        if text_hash != self._cached_text_hash:
            self._cached_text_hash = text_hash
            text_lower = text.lower()
            words = re.findall(r'\b[a-zA-Z]+\b', text_lower)
            self._cached_counts = Counter(words)

            # Extract and count phrases (2, 3, 4, 5-word ngrams)
            all_phrases: List[str] = []
            for n in range(2, 6):  # 2, 3, 4, 5-word phrases
                all_phrases.extend(self._extract_ngrams(words, n))

            self._cached_phrase_counts = Counter(all_phrases)

            # Build phrase positions for highlighting
            self._cached_phrase_positions.clear()
            for phrase, count in self._cached_phrase_counts.items():
                if count >= self.phrase_threshold and phrase not in self.IGNORE_PHRASES and phrase not in self._ignored_phrases:
                    # Find all positions of this phrase in the text
                    positions = []
                    start = 0
                    phrase_pattern = r'\b' + r'\s+'.join(re.escape(w) for w in phrase.split()) + r'\b'
                    for match in re.finditer(phrase_pattern, text_lower):
                        positions.append((match.start(), match.end()))
                    if positions:
                        self._cached_phrase_positions[phrase] = positions

    def analyze_text(self, text: str) -> Dict[str, int]:
        """Analyze text and return word frequencies for non-ignored words."""
        self.update_cache(text)

        # Filter out ignored words and short words
        return {
            word: count for word, count in self._cached_counts.items()
            if word not in self.IGNORE_WORDS
            and word not in self._ignored_words
            and len(word) > 2
            and count >= self.threshold
        }

    def is_overused(self, word: str, text: str) -> Tuple[bool, int]:
        """Check if a word is overused in the text.

        Returns:
            Tuple of (is_overused, count)
        """
        word_lower = word.lower()

        if word_lower in self.IGNORE_WORDS or word_lower in self._ignored_words:
            return False, 0

        if len(word_lower) <= 2:
            return False, 0

        # Use cached counts if available
        if self._cached_text_hash == hash(text):
            count = self._cached_counts.get(word_lower, 0)
        else:
            # Fallback to direct count (slower)
            self.update_cache(text)
            count = self._cached_counts.get(word_lower, 0)

        return count >= self.threshold, count

    def get_synonyms(self, word: str) -> List[str]:
        """Get synonyms for a word."""
        word_lower = word.lower()

        # Check our built-in synonyms first
        if word_lower in self.SYNONYMS:
            return self.SYNONYMS[word_lower]

        # Try nltk wordnet if available
        try:
            from nltk.corpus import wordnet
            synonyms = set()
            for syn in wordnet.synsets(word_lower):
                for lemma in syn.lemmas():
                    if lemma.name() != word_lower and '_' not in lemma.name():
                        synonyms.add(lemma.name())
            return list(synonyms)[:10]
        except:
            pass

        return []

    def ignore_word(self, word: str):
        """Add word to ignored list."""
        self._ignored_words.add(word.lower())

    def is_ignored(self, word: str) -> bool:
        """Check if word is in ignored list."""
        return word.lower() in self._ignored_words

    # Phrase detection methods

    def analyze_phrases(self, text: str) -> Dict[str, int]:
        """Analyze text and return phrase frequencies for non-ignored phrases.

        Returns phrases that appear at least phrase_threshold times.
        """
        self.update_cache(text)

        return {
            phrase: count for phrase, count in self._cached_phrase_counts.items()
            if phrase not in self.IGNORE_PHRASES
            and phrase not in self._ignored_phrases
            and count >= self.phrase_threshold
        }

    def is_phrase_overused(self, phrase: str) -> Tuple[bool, int]:
        """Check if a phrase is overused based on cached counts.

        Returns:
            Tuple of (is_overused, count)
        """
        phrase_lower = phrase.lower()

        if phrase_lower in self.IGNORE_PHRASES or phrase_lower in self._ignored_phrases:
            return False, 0

        count = self._cached_phrase_counts.get(phrase_lower, 0)
        return count >= self.phrase_threshold, count

    def get_overused_phrases_in_range(self, text: str, start: int, end: int) -> List[Tuple[str, int, int, int]]:
        """Get all overused phrases that overlap with a given text range.

        Args:
            text: Full document text (for cache validation)
            start: Start position in document
            end: End position in document

        Returns:
            List of (phrase, phrase_start, phrase_end, count) tuples
        """
        # Ensure cache is current
        if self._cached_text_hash != hash(text):
            self.update_cache(text)

        results = []
        for phrase, positions in self._cached_phrase_positions.items():
            count = self._cached_phrase_counts.get(phrase, 0)
            for pos_start, pos_end in positions:
                # Check if this phrase occurrence overlaps with our range
                if pos_start < end and pos_end > start:
                    results.append((phrase, pos_start, pos_end, count))

        return results

    def get_phrase_at_position(self, text: str, position: int) -> Optional[Tuple[str, int, int, int]]:
        """Get overused phrase at a specific position, if any.

        Args:
            text: Full document text
            position: Character position to check

        Returns:
            Tuple of (phrase, start, end, count) or None
        """
        # Ensure cache is current
        if self._cached_text_hash != hash(text):
            self.update_cache(text)

        # Check all overused phrases for one that contains this position
        # Prefer longer phrases over shorter ones
        best_match = None
        for phrase, positions in self._cached_phrase_positions.items():
            count = self._cached_phrase_counts.get(phrase, 0)
            for pos_start, pos_end in positions:
                if pos_start <= position < pos_end:
                    if best_match is None or len(phrase) > len(best_match[0]):
                        best_match = (phrase, pos_start, pos_end, count)

        return best_match

    def ignore_phrase(self, phrase: str):
        """Add phrase to ignored list."""
        self._ignored_phrases.add(phrase.lower())
        # Remove from cached positions
        phrase_lower = phrase.lower()
        if phrase_lower in self._cached_phrase_positions:
            del self._cached_phrase_positions[phrase_lower]

    def is_phrase_ignored(self, phrase: str) -> bool:
        """Check if phrase is in ignored list."""
        return phrase.lower() in self._ignored_phrases


@dataclass
class WritingStats:
    """Comprehensive writing statistics like ProWritingAid."""
    word_count: int = 0
    sentence_count: int = 0
    paragraph_count: int = 0
    avg_sentence_length: float = 0.0
    avg_word_length: float = 0.0

    # Sentence length variation
    short_sentences: int = 0  # < 10 words
    medium_sentences: int = 0  # 10-20 words
    long_sentences: int = 0  # 21-35 words
    very_long_sentences: int = 0  # > 35 words
    sentence_length_score: float = 0.0  # 0-100, higher is better variety

    # Sticky sentences (high % of glue words)
    sticky_sentences: List[str] = None
    sticky_sentence_count: int = 0
    avg_glue_index: float = 0.0

    # Echoes (repeated words in close proximity)
    echoes: List[Tuple[str, int, int]] = None  # (word, first_pos, second_pos)
    echo_count: int = 0

    # Passive voice
    passive_sentences: List[str] = None
    passive_count: int = 0
    passive_percentage: float = 0.0

    # Dialogue
    dialogue_percentage: float = 0.0
    dialogue_tags: Dict[str, int] = None  # tag -> count
    overused_dialogue_tags: List[str] = None

    # Readability
    flesch_reading_ease: float = 0.0
    flesch_grade_level: float = 0.0

    # Adverbs
    adverbs: List[str] = None
    adverb_count: int = 0
    adverb_percentage: float = 0.0

    # Clichés
    cliches: List[str] = None
    cliche_count: int = 0

    def __post_init__(self):
        """Initialize mutable defaults."""
        if self.sticky_sentences is None:
            self.sticky_sentences = []
        if self.echoes is None:
            self.echoes = []
        if self.passive_sentences is None:
            self.passive_sentences = []
        if self.dialogue_tags is None:
            self.dialogue_tags = {}
        if self.overused_dialogue_tags is None:
            self.overused_dialogue_tags = []
        if self.adverbs is None:
            self.adverbs = []
        if self.cliches is None:
            self.cliches = []


class ProWritingAnalyzer:
    """ProWritingAid-style comprehensive text analysis."""

    # Common glue words (function words that don't add meaning)
    GLUE_WORDS = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'been',
        'be', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
        'could', 'should', 'may', 'might', 'must', 'shall', 'can', 'need',
        'it', 'its', 'this', 'that', 'these', 'those', 'i', 'you', 'he',
        'she', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your',
        'his', 'our', 'their', 'what', 'which', 'who', 'whom', 'whose',
        'if', 'then', 'else', 'when', 'where', 'why', 'how', 'all', 'each',
        'every', 'both', 'few', 'more', 'most', 'other', 'some', 'such', 'no',
        'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 'just',
        'also', 'now', 'here', 'there', 'still', 'already', 'even', 'again',
        'into', 'through', 'during', 'before', 'after', 'above', 'below',
        'between', 'under', 'over', 'up', 'down', 'out', 'off', 'about',
        'any', 'am', 'being', 'having', 'doing', 'get', 'got', 'getting',
    }

    # Passive voice indicators
    PASSIVE_HELPERS = {'was', 'were', 'is', 'are', 'been', 'being', 'be', 'am'}

    # Common dialogue tags
    DIALOGUE_TAGS = {
        'said', 'asked', 'replied', 'answered', 'responded', 'exclaimed',
        'whispered', 'shouted', 'yelled', 'screamed', 'muttered', 'murmured',
        'declared', 'announced', 'stated', 'remarked', 'commented', 'noted',
        'added', 'continued', 'explained', 'insisted', 'demanded', 'suggested',
        'wondered', 'questioned', 'inquired', 'cried', 'called', 'sighed',
        'groaned', 'moaned', 'breathed', 'gasped', 'hissed', 'snapped',
        'barked', 'growled', 'laughed', 'chuckled', 'giggled', 'sobbed',
    }

    # Common clichés in writing
    CLICHES = {
        'at the end of the day', 'in the nick of time', 'all of a sudden',
        'in the blink of an eye', 'better late than never', 'as luck would have it',
        'a matter of time', 'the tip of the iceberg', 'a sigh of relief',
        'with bated breath', 'without a doubt', 'at first glance',
        'crystal clear', 'dead of night', 'in broad daylight', 'pitch black',
        'blood ran cold', 'heart skipped a beat', 'butterflies in stomach',
        'like a deer in headlights', 'like a fish out of water', 'cold as ice',
        'white as a sheet', 'red as a tomato', 'black as night', 'clear as day',
        'quiet as a mouse', 'sick as a dog', 'fit as a fiddle', 'stubborn as a mule',
        'light as a feather', 'heavy as lead', 'sharp as a tack', 'dumb as a rock',
        'needle in a haystack', 'easy as pie', 'piece of cake', 'walk in the park',
        'once upon a time', 'happily ever after', 'in the end', 'last but not least',
        'long story short', 'to make a long story short', 'bottom line',
        'at this point in time', 'going forward', 'on the same page',
        'think outside the box', 'low hanging fruit', 'push the envelope',
        'touch base', 'circle back', 'move the needle',
    }

    # Common adverbs to flag
    COMMON_ADVERBS = {
        'really', 'very', 'actually', 'basically', 'literally', 'totally',
        'absolutely', 'definitely', 'certainly', 'obviously', 'clearly',
        'simply', 'just', 'quite', 'rather', 'somewhat', 'fairly',
        'extremely', 'incredibly', 'particularly', 'especially',
        'suddenly', 'immediately', 'quickly', 'slowly', 'carefully',
        'quietly', 'loudly', 'softly', 'hardly', 'nearly', 'almost',
        'completely', 'entirely', 'perfectly', 'truly', 'honestly',
    }

    def __init__(self):
        """Initialize analyzer."""

    def analyze(self, text: str) -> WritingStats:
        """Perform comprehensive text analysis.

        Args:
            text: The text to analyze

        Returns:
            WritingStats with all metrics
        """
        stats = WritingStats()

        if not text or not text.strip():
            return stats

        # Basic counts
        words = re.findall(r'\b[a-zA-Z]+\b', text)
        stats.word_count = len(words)

        # Sentence detection
        sentences = self._split_sentences(text)
        stats.sentence_count = len(sentences)

        # Paragraph count
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        stats.paragraph_count = len(paragraphs) if paragraphs else 1

        if stats.word_count == 0 or stats.sentence_count == 0:
            return stats

        # Average lengths
        stats.avg_sentence_length = stats.word_count / stats.sentence_count
        stats.avg_word_length = sum(len(w) for w in words) / len(words)

        # Sentence length analysis
        self._analyze_sentence_lengths(sentences, stats)

        # Sticky sentences
        self._analyze_sticky_sentences(sentences, stats)

        # Echoes
        self._analyze_echoes(text, stats)

        # Passive voice
        self._analyze_passive_voice(sentences, stats)

        # Dialogue analysis
        self._analyze_dialogue(text, sentences, stats)

        # Readability
        self._calculate_readability(text, words, sentences, stats)

        # Adverbs
        self._analyze_adverbs(words, stats)

        # Clichés
        self._detect_cliches(text, stats)

        return stats

    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitter
        # Handle abbreviations, dialogue, etc.
        sentences = []

        # First, protect common abbreviations
        protected = text
        abbreviations = ['Mr.', 'Mrs.', 'Ms.', 'Dr.', 'Prof.', 'Jr.', 'Sr.',
                        'vs.', 'etc.', 'e.g.', 'i.e.', 'St.', 'Ave.', 'Blvd.']
        for abbr in abbreviations:
            protected = protected.replace(abbr, abbr.replace('.', '<DOT>'))

        # Split on sentence-ending punctuation
        parts = re.split(r'(?<=[.!?])\s+', protected)

        for part in parts:
            # Restore dots
            sentence = part.replace('<DOT>', '.')
            if sentence.strip():
                sentences.append(sentence.strip())

        return sentences

    def _analyze_sentence_lengths(self, sentences: List[str], stats: WritingStats):
        """Analyze sentence length distribution."""
        lengths = []

        for sentence in sentences:
            words = len(re.findall(r'\b[a-zA-Z]+\b', sentence))
            lengths.append(words)

            if words < 10:
                stats.short_sentences += 1
            elif words <= 20:
                stats.medium_sentences += 1
            elif words <= 35:
                stats.long_sentences += 1
            else:
                stats.very_long_sentences += 1

        # Calculate variety score (standard deviation normalized)
        if len(lengths) > 1:
            import statistics
            try:
                std_dev = statistics.stdev(lengths)
                mean = statistics.mean(lengths)
                # Coefficient of variation as variety indicator
                cv = (std_dev / mean * 100) if mean > 0 else 0
                # Score: higher CV = more variety = better (cap at 100)
                stats.sentence_length_score = min(100, cv * 2)
            except:
                stats.sentence_length_score = 50.0

    def _analyze_sticky_sentences(self, sentences: List[str], stats: WritingStats):
        """Analyze sentences for high glue word percentage (sticky sentences)."""
        sticky_threshold = 0.45  # 45% glue words = sticky
        total_glue_index = 0

        for sentence in sentences:
            words = re.findall(r'\b[a-zA-Z]+\b', sentence.lower())
            if len(words) < 4:
                continue

            glue_count = sum(1 for w in words if w in self.GLUE_WORDS)
            glue_index = glue_count / len(words) if words else 0
            total_glue_index += glue_index

            if glue_index >= sticky_threshold:
                stats.sticky_sentences.append(sentence[:100] + '...' if len(sentence) > 100 else sentence)

        stats.sticky_sentence_count = len(stats.sticky_sentences)
        stats.avg_glue_index = total_glue_index / len(sentences) if sentences else 0

    def _analyze_echoes(self, text: str, stats: WritingStats):
        """Find words repeated within close proximity (echoes)."""
        words = re.findall(r'\b([a-zA-Z]{4,})\b', text.lower())  # Only words 4+ chars
        word_positions: Dict[str, List[int]] = {}

        for i, word in enumerate(words):
            if word in self.GLUE_WORDS:
                continue
            if word not in word_positions:
                word_positions[word] = []
            word_positions[word].append(i)

        # Find echoes (same word within 10 words)
        echo_distance = 10
        for word, positions in word_positions.items():
            for i in range(len(positions) - 1):
                if positions[i+1] - positions[i] <= echo_distance:
                    stats.echoes.append((word, positions[i], positions[i+1]))

        stats.echo_count = len(stats.echoes)

    def _analyze_passive_voice(self, sentences: List[str], stats: WritingStats):
        """Detect passive voice constructions."""
        # Pattern: was/were/is/are/been/being + past participle (often ends in -ed)
        passive_pattern = re.compile(
            r'\b(was|were|is|are|am|been|being|be)\s+(?:\w+\s+)?(\w+ed|written|spoken|taken|given|made|done|seen|known|found|felt)\b',
            re.IGNORECASE
        )

        for sentence in sentences:
            if passive_pattern.search(sentence):
                stats.passive_sentences.append(sentence[:100] + '...' if len(sentence) > 100 else sentence)

        stats.passive_count = len(stats.passive_sentences)
        stats.passive_percentage = (stats.passive_count / len(sentences) * 100) if sentences else 0

    def _analyze_dialogue(self, text: str, sentences: List[str], stats: WritingStats):
        """Analyze dialogue and dialogue tags."""
        # Find dialogue (text in quotes)
        dialogue_matches = re.findall(r'["\u201c][^"\u201d]+["\u201d]', text)
        dialogue_words = sum(len(m.split()) for m in dialogue_matches)
        total_words = len(text.split())

        stats.dialogue_percentage = (dialogue_words / total_words * 100) if total_words > 0 else 0

        # Find dialogue tags
        tag_pattern = re.compile(
            r'["\u201d]\s*(?:,\s*)?(?:he|she|they|i|we|[A-Z][a-z]+)\s+(\w+ed?)',
            re.IGNORECASE
        )

        stats.dialogue_tags = Counter()
        for match in tag_pattern.finditer(text):
            tag = match.group(1).lower()
            if tag in self.DIALOGUE_TAGS:
                stats.dialogue_tags[tag] += 1

        # Find overused tags (>3 uses)
        stats.overused_dialogue_tags = [
            tag for tag, count in stats.dialogue_tags.items() if count > 3
        ]

    def _calculate_readability(self, text: str, words: List[str], sentences: List[str], stats: WritingStats):
        """Calculate Flesch reading ease and grade level."""
        if not words or not sentences:
            return

        # Count syllables (simple approximation)
        def count_syllables(word: str) -> int:
            word = word.lower()
            vowels = 'aeiouy'
            count = 0
            prev_vowel = False
            for char in word:
                is_vowel = char in vowels
                if is_vowel and not prev_vowel:
                    count += 1
                prev_vowel = is_vowel
            # Adjust for silent e
            if word.endswith('e') and count > 1:
                count -= 1
            return max(1, count)

        total_syllables = sum(count_syllables(w) for w in words)
        avg_syllables = total_syllables / len(words)
        avg_sentence_len = len(words) / len(sentences)

        # Flesch Reading Ease
        # 206.835 - 1.015 * (words/sentences) - 84.6 * (syllables/words)
        stats.flesch_reading_ease = 206.835 - (1.015 * avg_sentence_len) - (84.6 * avg_syllables)
        stats.flesch_reading_ease = max(0, min(100, stats.flesch_reading_ease))

        # Flesch-Kincaid Grade Level
        # 0.39 * (words/sentences) + 11.8 * (syllables/words) - 15.59
        stats.flesch_grade_level = (0.39 * avg_sentence_len) + (11.8 * avg_syllables) - 15.59
        stats.flesch_grade_level = max(0, stats.flesch_grade_level)

    def _analyze_adverbs(self, words: List[str], stats: WritingStats):
        """Analyze adverb usage."""
        word_list = [w.lower() for w in words]

        # Find -ly adverbs and common adverbs
        for word in word_list:
            if word.endswith('ly') and len(word) > 3:
                if word not in stats.adverbs:
                    stats.adverbs.append(word)
            elif word in self.COMMON_ADVERBS:
                if word not in stats.adverbs:
                    stats.adverbs.append(word)

        stats.adverb_count = sum(1 for w in word_list if w in stats.adverbs or (w.endswith('ly') and len(w) > 3))
        stats.adverb_percentage = (stats.adverb_count / len(words) * 100) if words else 0

    def _detect_cliches(self, text: str, stats: WritingStats):
        """Detect clichés in text."""
        text_lower = text.lower()

        for cliche in self.CLICHES:
            if cliche in text_lower:
                stats.cliches.append(cliche)

        stats.cliche_count = len(stats.cliches)

    def format_report(self, stats: WritingStats) -> str:
        """Format stats as an HTML report."""
        html = """
        <style>
            .stat-section { margin: 15px 0; padding: 10px; background: #f5f5f5; border-radius: 4px; }
            .stat-header { font-weight: bold; color: #374151; margin-bottom: 8px; }
            .stat-value { color: #059669; font-weight: bold; }
            .stat-warning { color: #dc2626; }
            .stat-ok { color: #059669; }
            .meter { height: 10px; background: #e5e7eb; border-radius: 5px; margin: 5px 0; }
            .meter-fill { height: 100%; border-radius: 5px; }
            .list-item { padding: 4px 0; border-bottom: 1px solid #e5e7eb; }
        </style>
        """

        # Overview
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>📊 Document Overview</div>
            <p>Words: <span class='stat-value'>{stats.word_count:,}</span> |
               Sentences: <span class='stat-value'>{stats.sentence_count:,}</span> |
               Paragraphs: <span class='stat-value'>{stats.paragraph_count:,}</span></p>
            <p>Avg sentence length: <span class='stat-value'>{stats.avg_sentence_length:.1f}</span> words |
               Avg word length: <span class='stat-value'>{stats.avg_word_length:.1f}</span> chars</p>
        </div>
        """

        # Sentence Length Variation
        variety_class = 'stat-ok' if stats.sentence_length_score > 40 else 'stat-warning'
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>📏 Sentence Length Variation</div>
            <p>Short (< 10 words): {stats.short_sentences} |
               Medium (10-20): {stats.medium_sentences} |
               Long (21-35): {stats.long_sentences} |
               Very Long (> 35): {stats.very_long_sentences}</p>
            <p>Variety Score: <span class='{variety_class}'>{stats.sentence_length_score:.0f}/100</span></p>
            <div class='meter'><div class='meter-fill' style='width: {stats.sentence_length_score}%; background: {"#059669" if stats.sentence_length_score > 40 else "#dc2626"};'></div></div>
        </div>
        """

        # Sticky Sentences
        sticky_class = 'stat-ok' if stats.sticky_sentence_count < 3 else 'stat-warning'
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>🔗 Sticky Sentences (high glue word %)</div>
            <p>Sticky sentences: <span class='{sticky_class}'>{stats.sticky_sentence_count}</span> |
               Avg glue index: {stats.avg_glue_index:.1%}</p>
        """
        if stats.sticky_sentences[:3]:
            html += "<div style='font-size: 11px; color: #666; margin-top: 5px;'>"
            for s in stats.sticky_sentences[:3]:
                html += f"<div class='list-item'>• {s}</div>"
            if len(stats.sticky_sentences) > 3:
                html += f"<div class='list-item'><em>...and {len(stats.sticky_sentences) - 3} more</em></div>"
            html += "</div>"
        html += "</div>"

        # Echoes
        echo_class = 'stat-ok' if stats.echo_count < 5 else 'stat-warning'
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>🔁 Echoes (repeated words nearby)</div>
            <p>Echoes found: <span class='{echo_class}'>{stats.echo_count}</span></p>
        """
        if stats.echoes[:5]:
            echo_words = list(set([e[0] for e in stats.echoes[:10]]))
            html += f"<p style='font-size: 11px; color: #666;'>Words: {', '.join(echo_words[:8])}</p>"
        html += "</div>"

        # Passive Voice
        passive_class = 'stat-ok' if stats.passive_percentage < 15 else 'stat-warning'
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>🔄 Passive Voice</div>
            <p>Passive sentences: <span class='{passive_class}'>{stats.passive_count}</span> ({stats.passive_percentage:.1f}%)</p>
        </div>
        """

        # Dialogue
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>💬 Dialogue Analysis</div>
            <p>Dialogue: {stats.dialogue_percentage:.1f}% of text</p>
        """
        if stats.dialogue_tags:
            top_tags = sorted(stats.dialogue_tags.items(), key=lambda x: -x[1])[:5]
            html += "<p style='font-size: 11px;'>Tags: "
            html += ', '.join([f"{tag} ({count})" for tag, count in top_tags])
            html += "</p>"
        if stats.overused_dialogue_tags:
            html += f"<p class='stat-warning' style='font-size: 11px;'>Overused: {', '.join(stats.overused_dialogue_tags)}</p>"
        html += "</div>"

        # Readability
        ease_desc = self._get_readability_description(stats.flesch_reading_ease)
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>📖 Readability</div>
            <p>Flesch Reading Ease: <span class='stat-value'>{stats.flesch_reading_ease:.1f}</span> ({ease_desc})</p>
            <p>Grade Level: <span class='stat-value'>{stats.flesch_grade_level:.1f}</span></p>
        </div>
        """

        # Adverbs
        adverb_class = 'stat-ok' if stats.adverb_percentage < 3 else 'stat-warning'
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>📝 Adverb Usage</div>
            <p>Adverbs: <span class='{adverb_class}'>{stats.adverb_count}</span> ({stats.adverb_percentage:.1f}%)</p>
        """
        if stats.adverbs[:10]:
            html += f"<p style='font-size: 11px; color: #666;'>Examples: {', '.join(stats.adverbs[:10])}</p>"
        html += "</div>"

        # Clichés
        cliche_class = 'stat-ok' if stats.cliche_count == 0 else 'stat-warning'
        html += f"""
        <div class='stat-section'>
            <div class='stat-header'>⚠️ Clichés</div>
            <p>Clichés found: <span class='{cliche_class}'>{stats.cliche_count}</span></p>
        """
        if stats.cliches:
            html += f"<p style='font-size: 11px; color: #666;'>{', '.join(stats.cliches[:5])}</p>"
        html += "</div>"

        return html

    def _get_readability_description(self, score: float) -> str:
        """Get description for Flesch reading ease score."""
        if score >= 90:
            return "Very Easy"
        elif score >= 80:
            return "Easy"
        elif score >= 70:
            return "Fairly Easy"
        elif score >= 60:
            return "Standard"
        elif score >= 50:
            return "Fairly Difficult"
        elif score >= 30:
            return "Difficult"
        else:
            return "Very Difficult"


class GrammarChecker:
    """Basic grammar checker using language-tool-python, gingerit, or basic rules."""

    def __init__(self):
        """Initialize grammar checker."""
        self._tool = None
        self._use_language_tool = False
        self._use_gingerit = False
        self._ignored_rules: Set[str] = set()
        # Don't auto-init language tool - it's slow and resource intensive
        # Only use basic rules for real-time checking
        self._language_tool_available = self._check_language_tool_available()
        self._gingerit_available = self._check_gingerit_available()

        # Log available grammar tools
        if self._language_tool_available:
            print("Grammar checker: language-tool-python available (requires Java)")
        elif self._gingerit_available:
            print("Grammar checker: Using gingerit as fallback (no Java required)")
        else:
            print("Grammar checker: Using basic rules only")

    def _check_java_available(self) -> bool:
        """Check if Java is available on the system."""
        import subprocess
        import shutil

        # First check if java is in PATH
        java_path = shutil.which('java')
        if not java_path:
            return False

        # Try to run java -version
        try:
            result = subprocess.run(
                [java_path, '-version'],
                capture_output=True,
                timeout=5
            )
            return result.returncode == 0
        except Exception:
            return False

    def _check_language_tool_available(self) -> bool:
        """Check if language-tool-python is available (requires Java)."""
        try:
            pass
            # Also check if Java is available
            if not self._check_java_available():
                print("Grammar checker: language-tool-python installed but Java not available")
                return False
            return True
        except ImportError:
            return False

    def _check_gingerit_available(self) -> bool:
        """Check if gingerit is available as a fallback."""
        try:
            return True
        except ImportError:
            return False

    def _init_language_tool(self):
        """Initialize the language tool (only when explicitly needed)."""
        if self._tool is not None:
            return

        # Try language-tool-python first (requires Java)
        if self._language_tool_available:
            try:
                import language_tool_python
                self._tool = language_tool_python.LanguageTool('en-US')
                self._use_language_tool = True
                print("Grammar checker: Initialized language-tool-python")
                return
            except Exception as e:
                print(f"Grammar checker: language-tool-python error ({e})")
                self._language_tool_available = False

        # Fall back to gingerit (no Java required, uses web API)
        if self._gingerit_available:
            try:
                from gingerit.gingerit import GingerIt
                self._tool = GingerIt()
                self._use_gingerit = True
                print("Grammar checker: Initialized gingerit (fallback)")
                return
            except Exception as e:
                print(f"Grammar checker: gingerit error ({e})")
                self._gingerit_available = False

        # No advanced grammar tool available
        print("Grammar checker: Using basic rules only (no advanced tool available)")
        self._use_language_tool = False
        self._use_gingerit = False

    def check_text(self, text: str, use_full_check: bool = False) -> List[Tuple[int, int, str, str, List[str]]]:
        """Check text for grammar errors.

        Args:
            text: Text to check
            use_full_check: If True, use heavy language-tool or gingerit (only for explicit recheck)

        Returns:
            List of (start, end, error_text, message, suggestions)
        """
        errors = []

        # Only use advanced tools for explicit full checks (recheck button)
        if use_full_check and (self._language_tool_available or self._gingerit_available):
            self._init_language_tool()

            # Try language-tool-python first
            if self._use_language_tool and self._tool:
                try:
                    matches = self._tool.check(text)
                    for match in matches:
                        rule = getattr(match, 'rule_id', None) or getattr(match, 'ruleId', '')
                        err_len = getattr(match, 'error_length', None) or getattr(match, 'errorLength', 0)
                        if rule not in self._ignored_rules:
                            errors.append((
                                match.offset,
                                match.offset + err_len,
                                text[match.offset:match.offset + err_len],
                                match.message,
                                match.replacements[:5] if match.replacements else []
                            ))
                    return errors
                except Exception as e:
                    print(f"Grammar check error (language-tool): {e}")

            # Fall back to gingerit
            if self._use_gingerit and self._tool:
                try:
                    result = self._tool.parse(text)
                    corrections = result.get('corrections', [])
                    for correction in corrections:
                        start = correction.get('start', 0)
                        end = start + len(correction.get('text', ''))
                        errors.append((
                            start,
                            end,
                            correction.get('text', ''),
                            correction.get('definition', 'Grammar suggestion'),
                            [correction.get('correct', '')] if correction.get('correct') else []
                        ))
                    return errors
                except Exception as e:
                    print(f"Grammar check error (gingerit): {e}")

        # For real-time checking, use fast basic rules only
        errors.extend(self._check_basic_rules(text))
        return errors

    def _check_basic_rules(self, text: str) -> List[Tuple[int, int, str, str, List[str]]]:
        """Apply basic grammar rules."""
        errors = []

        # Double word detection
        double_word_pattern = r'\b(\w+)\s+\1\b'
        for match in re.finditer(double_word_pattern, text, re.IGNORECASE):
            word = match.group(1)
            errors.append((
                match.start(),
                match.end(),
                match.group(),
                f"Repeated word: '{word}'",
                [word]
            ))

        # a/an usage
        a_before_vowel = r'\ba\s+([aeiouAEIOU]\w*)\b'
        for match in re.finditer(a_before_vowel, text):
            errors.append((
                match.start(),
                match.start() + 1,
                'a',
                f"Use 'an' before words starting with a vowel sound",
                ['an']
            ))

        an_before_consonant = r'\ban\s+([bcdfghjklmnpqrstvwxyzBCDFGHJKLMNPQRSTVWXYZ]\w*)\b'
        for match in re.finditer(an_before_consonant, text):
            # Skip words that sound like they start with vowels (hour, honest, etc.)
            word = match.group(1).lower()
            if word not in ['hour', 'honest', 'honor', 'honour', 'heir', 'herb']:
                errors.append((
                    match.start(),
                    match.start() + 2,
                    'an',
                    f"Use 'a' before words starting with a consonant sound",
                    ['a']
                ))

        return errors

    def ignore_rule(self, rule_id: str):
        """Ignore a specific grammar rule."""
        self._ignored_rules.add(rule_id)

    def is_available(self) -> bool:
        """Check if grammar checking is available."""
        return self._use_language_tool or True  # Basic rules always available


class WritingHighlighter(QSyntaxHighlighter):
    """Syntax highlighter for spelling, grammar, and overuse checking."""

    def __init__(
        self,
        document: QTextDocument,
        spell_checker: SpellChecker,
        grammar_checker: Optional[GrammarChecker] = None,
        overuse_detector: Optional[OveruseDetector] = None
    ):
        super().__init__(document)
        self.spell_checker = spell_checker
        self.grammar_checker = grammar_checker or GrammarChecker()
        self.overuse_detector = overuse_detector or OveruseDetector()

        # Check modes: ON_DEMAND (green), REALTIME (blue), OFF (white)
        # Start in ON_DEMAND mode - initial check only, then on recheck button
        self._spell_mode = CheckMode.ON_DEMAND
        self._grammar_mode = CheckMode.ON_DEMAND
        self._overuse_mode = CheckMode.ON_DEMAND

        # Flag to indicate we're doing an on-demand check (recheck button or initial)
        self._doing_on_demand_check = False

        # Full grammar check mode (for explicit recheck only)
        self._full_grammar_check = False

        # Store errors by block: block -> [WritingError]
        self._errors: Dict[int, List[WritingError]] = {}

        # Full document text for overuse analysis
        self._full_text = ""

        # Ignored errors by word only (more robust to text changes)
        # Format: word -> error_type (so we can ignore specific types)
        self._ignored_words: Dict[str, Set[ErrorType]] = {}

        # Track block content hashes to detect changes
        self._block_hashes: Dict[int, int] = {}

        # Connect to document changes to clear stale data
        document.contentsChange.connect(self._on_contents_change)

        # Spelling format (red wavy underline)
        self.spelling_format = QTextCharFormat()
        self.spelling_format.setUnderlineStyle(
            QTextCharFormat.UnderlineStyle.SpellCheckUnderline
        )
        self.spelling_format.setUnderlineColor(QColor(255, 0, 0))  # Red

        # Grammar format (green wavy underline)
        self.grammar_format = QTextCharFormat()
        self.grammar_format.setUnderlineStyle(
            QTextCharFormat.UnderlineStyle.SpellCheckUnderline
        )
        self.grammar_format.setUnderlineColor(QColor(0, 180, 0))  # Green

        # Overuse format (blue wavy underline)
        self.overuse_format = QTextCharFormat()
        self.overuse_format.setUnderlineStyle(
            QTextCharFormat.UnderlineStyle.SpellCheckUnderline
        )
        self.overuse_format.setUnderlineColor(QColor(0, 120, 255))  # Blue

        # Markdown heading formats
        self.md_title_format = QTextCharFormat()
        self.md_title_format.setFontPointSize(28)
        self.md_title_format.setFontWeight(QFont.Weight.Bold)

        self.md_h1_format = QTextCharFormat()
        self.md_h1_format.setFontPointSize(24)
        self.md_h1_format.setFontWeight(QFont.Weight.Bold)

        self.md_h2_format = QTextCharFormat()
        self.md_h2_format.setFontPointSize(18)
        self.md_h2_format.setFontWeight(QFont.Weight.Bold)

        self.md_h3_format = QTextCharFormat()
        self.md_h3_format.setFontPointSize(14)
        self.md_h3_format.setFontWeight(QFont.Weight.Bold)

        self.md_h4_format = QTextCharFormat()
        self.md_h4_format.setFontPointSize(12)
        self.md_h4_format.setFontWeight(QFont.Weight.Bold)

        # Markdown heading prefix format (hidden - very small and transparent)
        self.md_prefix_format = QTextCharFormat()
        self.md_prefix_format.setFontPointSize(1)  # Tiny font
        self.md_prefix_format.setForeground(QColor(255, 255, 255, 0))  # Fully transparent

        # Markdown bold format
        self.md_bold_format = QTextCharFormat()
        self.md_bold_format.setFontWeight(QFont.Weight.Bold)

        # Markdown italic format
        self.md_italic_format = QTextCharFormat()
        self.md_italic_format.setFontItalic(True)

        # Markdown bold+italic format
        self.md_bold_italic_format = QTextCharFormat()
        self.md_bold_italic_format.setFontWeight(QFont.Weight.Bold)
        self.md_bold_italic_format.setFontItalic(True)

        # Markdown marker format (hidden - very small and transparent)
        self.md_marker_format = QTextCharFormat()
        self.md_marker_format.setFontPointSize(1)  # Tiny font
        self.md_marker_format.setForeground(QColor(255, 255, 255, 0))  # Fully transparent

    def _on_contents_change(self, position: int, chars_removed: int, chars_added: int):
        """Handle document content changes to clear stale error data."""
        if chars_removed > 0 or chars_added > 0:
            # Find affected blocks
            doc = self.document()
            if not doc:
                return

            # Get block at position
            block = doc.findBlock(position)
            if not block.isValid():
                return

            # Clear errors for affected blocks and following blocks if content shifted
            start_block = block.blockNumber()

            # If chars were added or removed, block numbers after this point may have shifted
            # Clear all error data from this block onward to be safe
            blocks_to_clear = [bn for bn in self._errors.keys() if bn >= start_block]
            for bn in blocks_to_clear:
                if bn in self._errors:
                    del self._errors[bn]
                if bn in self._block_hashes:
                    del self._block_hashes[bn]

    def set_spell_mode(self, mode: CheckMode):
        """Set spell checking mode."""
        self._spell_mode = mode
        # Only rehighlight if switching to realtime (other modes handled by recheck)
        if mode == CheckMode.REALTIME:
            self.rehighlight()

    def set_grammar_mode(self, mode: CheckMode):
        """Set grammar checking mode."""
        self._grammar_mode = mode
        if mode == CheckMode.REALTIME:
            self.rehighlight()

    def set_overuse_mode(self, mode: CheckMode):
        """Set overuse detection mode."""
        self._overuse_mode = mode
        if mode == CheckMode.REALTIME:
            self.rehighlight()

    def get_spell_mode(self) -> CheckMode:
        """Get current spell checking mode."""
        return self._spell_mode

    def get_grammar_mode(self) -> CheckMode:
        """Get current grammar checking mode."""
        return self._grammar_mode

    def get_overuse_mode(self) -> CheckMode:
        """Get current overuse detection mode."""
        return self._overuse_mode

    def update_full_text(self, text: str):
        """Update the full document text for overuse analysis."""
        self._full_text = text

    def do_full_recheck(self):
        """Perform a full recheck including heavy grammar analysis (on-demand checks)."""
        self._full_grammar_check = True
        self._doing_on_demand_check = True
        self.rehighlight()
        self._full_grammar_check = False
        self._doing_on_demand_check = False

    def do_initial_check(self):
        """Perform initial check when document is loaded."""
        self._doing_on_demand_check = True
        self.rehighlight()
        self._doing_on_demand_check = False

    def ignore_error(self, block_number: int, start: int, word: str, error_type: Optional[ErrorType] = None):
        """Ignore a specific error.

        Args:
            block_number: Block number (kept for API compatibility, but not used)
            start: Start position (kept for API compatibility, but not used)
            word: The word to ignore
            error_type: Optional specific error type to ignore for this word
        """
        word_lower = word.lower()
        if word_lower not in self._ignored_words:
            self._ignored_words[word_lower] = set()

        if error_type:
            # Ignore only this specific error type for this word
            self._ignored_words[word_lower].add(error_type)
        else:
            # Ignore all error types for this word
            self._ignored_words[word_lower].update([ErrorType.SPELLING, ErrorType.GRAMMAR, ErrorType.OVERUSE, ErrorType.OVERUSED_PHRASE])

        self.rehighlight()

    def is_word_ignored(self, word: str, error_type: ErrorType) -> bool:
        """Check if a word is ignored for a specific error type."""
        word_lower = word.lower()
        if word_lower in self._ignored_words:
            return error_type in self._ignored_words[word_lower]
        return False

    def clear_ignored(self):
        """Clear all ignored words."""
        self._ignored_words.clear()
        self.rehighlight()

    def get_error_at(self, block_number: int, position_in_block: int) -> Optional[WritingError]:
        """Get error at position."""
        if block_number not in self._errors:
            return None

        for error in self._errors[block_number]:
            if error.start <= position_in_block < error.end:
                return error
        return None

    def _should_check(self, mode: CheckMode) -> bool:
        """Determine if a check should run based on mode and current state."""
        if mode == CheckMode.OFF:
            return False
        if mode == CheckMode.REALTIME:
            return True
        # ON_DEMAND: only check during on-demand operations (recheck button or initial)
        return self._doing_on_demand_check

    def _apply_markdown_formatting(self, text: str):
        """Apply visual formatting for Markdown syntax.

        This renders Markdown headings, bold, and italic with appropriate fonts
        while keeping the Markdown syntax visible (but dimmed).
        """
        # Check for heading prefixes at start of line
        stripped = text.lstrip()
        leading_spaces = len(text) - len(stripped)

        # Heading detection (most specific first)
        if stripped.startswith("##### "):
            # Heading 4
            prefix_len = 6 + leading_spaces
            self.setFormat(0, prefix_len, self.md_prefix_format)
            self.setFormat(prefix_len, len(text) - prefix_len, self.md_h4_format)
        elif stripped.startswith("#### "):
            # Heading 3
            prefix_len = 5 + leading_spaces
            self.setFormat(0, prefix_len, self.md_prefix_format)
            self.setFormat(prefix_len, len(text) - prefix_len, self.md_h3_format)
        elif stripped.startswith("### "):
            # Heading 2
            prefix_len = 4 + leading_spaces
            self.setFormat(0, prefix_len, self.md_prefix_format)
            self.setFormat(prefix_len, len(text) - prefix_len, self.md_h2_format)
        elif stripped.startswith("## "):
            # Heading 1
            prefix_len = 3 + leading_spaces
            self.setFormat(0, prefix_len, self.md_prefix_format)
            self.setFormat(prefix_len, len(text) - prefix_len, self.md_h1_format)
        elif stripped.startswith("# "):
            # Title
            prefix_len = 2 + leading_spaces
            self.setFormat(0, prefix_len, self.md_prefix_format)
            self.setFormat(prefix_len, len(text) - prefix_len, self.md_title_format)

        # Bold-italic (***text***) - must come before bold and italic
        for match in re.finditer(r'\*\*\*(.+?)\*\*\*', text):
            # Dim the markers
            self.setFormat(match.start(), 3, self.md_marker_format)
            self.setFormat(match.end() - 3, 3, self.md_marker_format)
            # Bold+italic the content
            self.setFormat(match.start() + 3, len(match.group(1)), self.md_bold_italic_format)

        # Bold (**text**)
        for match in re.finditer(r'\*\*(?!\*)(.+?)(?<!\*)\*\*', text):
            # Dim the markers
            self.setFormat(match.start(), 2, self.md_marker_format)
            self.setFormat(match.end() - 2, 2, self.md_marker_format)
            # Bold the content
            self.setFormat(match.start() + 2, len(match.group(1)), self.md_bold_format)

        # Italic (*text* but not **text**)
        for match in re.finditer(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', text):
            # Dim the markers
            self.setFormat(match.start(), 1, self.md_marker_format)
            self.setFormat(match.end() - 1, 1, self.md_marker_format)
            # Italic the content
            self.setFormat(match.start() + 1, len(match.group(1)), self.md_italic_format)

        # Italic with underscores (_text_)
        for match in re.finditer(r'(?<![_\w])_([^_]+)_(?![_\w])', text):
            # Dim the markers
            self.setFormat(match.start(), 1, self.md_marker_format)
            self.setFormat(match.end() - 1, 1, self.md_marker_format)
            # Italic the content
            self.setFormat(match.start() + 1, len(match.group(1)), self.md_italic_format)

    def highlightBlock(self, text: str):
        """Highlight errors in a block of text."""
        block = self.currentBlock()
        block_number = block.blockNumber()

        # Store current block hash to track changes
        current_hash = hash(text)
        self._block_hashes[block_number] = current_hash

        # Clear previous errors for this block
        self._errors[block_number] = []

        # Apply Markdown formatting first (before error checks)
        self._apply_markdown_formatting(text)

        # Helper to get the existing character format at a position in the block
        # This preserves font size, weight, etc. when we add underlines
        def get_existing_format(pos: int) -> QTextCharFormat:
            """Get the document's character format at position within current block."""
            cursor = QTextCursor(block)
            cursor.setPosition(block.position() + pos)
            return cursor.charFormat()

        # Spell checking - respects mode
        if self._should_check(self._spell_mode):
            for match in re.finditer(r'\b[a-zA-Z]+\b', text):
                word = match.group()
                # Skip if ignored
                if self.is_word_ignored(word, ErrorType.SPELLING):
                    continue

                is_correct, reason = self.spell_checker.check_word(word)
                if not is_correct:
                    # Get existing format and merge with our underline format
                    # This preserves font size, weight, etc.
                    existing_format = get_existing_format(match.start())
                    existing_format.merge(self.spelling_format)
                    self.setFormat(match.start(), len(word), existing_format)
                    suggestions = self.spell_checker.suggest(word)
                    self._errors[block_number].append(WritingError(
                        error_type=ErrorType.SPELLING,
                        word=word,
                        start=match.start(),
                        end=match.end(),
                        reason=reason,
                        suggestions=suggestions
                    ))

        # Grammar checking - respects mode
        # Use full check only during explicit recheck, otherwise use fast basic rules
        if self._should_check(self._grammar_mode) and self.grammar_checker:
            grammar_errors = self.grammar_checker.check_text(text, use_full_check=self._full_grammar_check)
            for start, end, error_text, message, suggestions in grammar_errors:
                # Skip if ignored
                if self.is_word_ignored(error_text, ErrorType.GRAMMAR):
                    continue

                # Don't double-flag spelling errors
                already_flagged = any(
                    e.start == start and e.end == end
                    for e in self._errors[block_number]
                )
                if not already_flagged:
                    # Merge with existing format to preserve font size, weight, etc.
                    existing_format = get_existing_format(start)
                    existing_format.merge(self.grammar_format)
                    self.setFormat(start, end - start, existing_format)
                    self._errors[block_number].append(WritingError(
                        error_type=ErrorType.GRAMMAR,
                        word=error_text,
                        start=start,
                        end=end,
                        reason=message,
                        suggestions=suggestions
                    ))

        # Overuse detection - respects mode
        if self._should_check(self._overuse_mode) and self._full_text and self.overuse_detector:
            for match in re.finditer(r'\b[a-zA-Z]+\b', text):
                word = match.group()
                # Skip if ignored
                if self.is_word_ignored(word, ErrorType.OVERUSE):
                    continue

                is_overused, count = self.overuse_detector.is_overused(word, self._full_text)
                if is_overused:
                    # Don't double-flag
                    already_flagged = any(
                        e.start == match.start() and e.end == match.end()
                        for e in self._errors[block_number]
                    )
                    if not already_flagged:
                        # Merge with existing format to preserve font size, weight, etc.
                        existing_format = get_existing_format(match.start())
                        existing_format.merge(self.overuse_format)
                        self.setFormat(match.start(), len(word), existing_format)
                        synonyms = self.overuse_detector.get_synonyms(word)
                        self._errors[block_number].append(WritingError(
                            error_type=ErrorType.OVERUSE,
                            word=word,
                            start=match.start(),
                            end=match.end(),
                            reason=f"Word '{word}' appears {count} times. Consider varying your word choice.",
                            suggestions=synonyms
                        ))

            # Overused phrase detection (2, 3, 4, 5-word phrases)
            # Calculate block start position in full document
            block_start = block.position()
            block_end = block_start + len(text)

            # Get overused phrases that overlap with this block
            phrase_occurrences = self.overuse_detector.get_overused_phrases_in_range(
                self._full_text, block_start, block_end
            )

            for phrase, phrase_start, phrase_end, count in phrase_occurrences:
                # Convert document position to block position
                local_start = max(0, phrase_start - block_start)
                local_end = min(len(text), phrase_end - block_start)

                # Skip if the phrase doesn't actually overlap this block's text range
                if local_start >= len(text) or local_end <= 0:
                    continue

                # Skip if ignored
                if self.is_word_ignored(phrase, ErrorType.OVERUSED_PHRASE):
                    continue

                # Don't double-flag if already flagged
                already_flagged = any(
                    e.start == local_start and e.end == local_end
                    for e in self._errors[block_number]
                )
                if not already_flagged:
                    # Use a different shade of blue for phrases (purple-ish)
                    phrase_format = QTextCharFormat()
                    phrase_format.setUnderlineStyle(
                        QTextCharFormat.UnderlineStyle.SpellCheckUnderline
                    )
                    phrase_format.setUnderlineColor(QColor(148, 0, 211))  # Purple for phrases

                    existing_format = get_existing_format(local_start)
                    existing_format.merge(phrase_format)
                    self.setFormat(local_start, local_end - local_start, existing_format)

                    word_count = len(phrase.split())
                    self._errors[block_number].append(WritingError(
                        error_type=ErrorType.OVERUSED_PHRASE,
                        word=phrase,
                        start=local_start,
                        end=local_end,
                        reason=f"Phrase '{phrase}' ({word_count} words) appears {count} times. Consider rephrasing.",
                        suggestions=[]  # No automatic suggestions for phrases
                    ))


class ContextLookupDialog(QDialog):
    """Dialog for displaying context lookup results."""

    def __init__(self, title: str, content: str, parent=None):
        """Initialize context lookup dialog."""
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setMinimumSize(600, 400)

        layout = QVBoxLayout(self)

        # Content display
        self.content_browser = QTextBrowser()
        self.content_browser.setMarkdown(content)
        self.content_browser.setOpenExternalLinks(False)
        layout.addWidget(self.content_browser)

        # Close button
        close_button = QPushButton("Close")
        close_button.clicked.connect(self.accept)
        layout.addWidget(close_button)


class QuickReferenceDialog(QDialog):
    """Dialog for quick reference selection."""

    def __init__(self, items: list, item_type: str, parent=None):
        """Initialize quick reference dialog."""
        super().__init__(parent)
        self.setWindowTitle(f"Select {item_type}")
        self.setMinimumSize(400, 300)
        self.selected_item = None

        layout = QVBoxLayout(self)

        label = QLabel(f"Select a {item_type} to view:")
        layout.addWidget(label)

        # Item list
        self.item_list = QListWidget()
        for item in items:
            self.item_list.addItem(item)
        self.item_list.itemDoubleClicked.connect(self._on_item_selected)
        layout.addWidget(self.item_list)

        # Buttons
        button_layout = QHBoxLayout()

        select_button = QPushButton("Select")
        select_button.clicked.connect(self._on_item_selected)
        button_layout.addWidget(select_button)

        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        button_layout.addWidget(cancel_button)

        layout.addLayout(button_layout)

    def _on_item_selected(self):
        """Handle item selection."""
        current_item = self.item_list.currentItem()
        if current_item:
            self.selected_item = current_item.text()
            self.accept()


class EnhancedTextEditor(QTextEdit):
    """Enhanced text editor with context menu, headers, spell/grammar/overuse check, Word support, and TTS."""

    context_lookup_requested = pyqtSignal(str)
    spell_check_toggled = pyqtSignal(bool)
    grammar_check_toggled = pyqtSignal(bool)
    overuse_check_toggled = pyqtSignal(bool)

    # TTS signals
    tts_started = pyqtSignal()
    tts_stopped = pyqtSignal()
    tts_error = pyqtSignal(str)
    tts_progress = pyqtSignal(str)  # Progress updates for VibeVoice

    # Heading style definitions for display and export
    # Format: (font_size, bold, top_margin, bottom_margin, heading_level)
    HEADING_STYLES = {
        "Normal": (12, False, 0, 0, HeadingLevel.NORMAL),
        "Title": (28, True, 24, 12, HeadingLevel.TITLE),
        "Heading 1": (24, True, 18, 8, HeadingLevel.HEADING_1),
        "Heading 2": (18, True, 14, 6, HeadingLevel.HEADING_2),
        "Heading 3": (14, True, 10, 4, HeadingLevel.HEADING_3),
        "Heading 4": (12, True, 8, 4, HeadingLevel.HEADING_4),
    }

    # Legacy mapping for backward compatibility
    HEADER_SIZES = {
        "Normal": 12,
        "Title": 28,
        "Heading 1": 24,
        "Heading 2": 18,
        "Heading 3": 14,
        "Heading 4": 12,
    }

    def __init__(self, parent=None):
        """Initialize enhanced text editor."""
        super().__init__(parent)

        # Callbacks for context lookup
        self.lookup_worldbuilding_callback: Optional[Callable] = None
        self.lookup_characters_callback: Optional[Callable] = None
        self.lookup_plot_callback: Optional[Callable] = None
        self.lookup_context_callback: Optional[Callable] = None
        self.get_character_list_callback: Optional[Callable] = None
        self.get_worldbuilding_sections_callback: Optional[Callable] = None

        # Enable rich text
        self.setAcceptRichText(True)

        # Initialize checkers
        self.spell_checker = SpellChecker()
        self.grammar_checker = GrammarChecker()
        self.overuse_detector = OveruseDetector()

        # Initialize combined highlighter
        self.writing_highlighter = WritingHighlighter(
            self.document(),
            self.spell_checker,
            self.grammar_checker,
            self.overuse_detector
        )

        # Enable mouse tracking for tooltips on both widget and viewport
        # On macOS, QTextEdit uses a viewport for rendering, so we must enable
        # mouse tracking there and handle tooltip events via viewportEvent()
        self.setMouseTracking(True)
        self.viewport().setMouseTracking(True)

        # Update overuse detection when text changes (with longer debounce for performance)
        self.textChanged.connect(self._on_text_changed)
        self._update_timer = QTimer()
        self._update_timer.setSingleShot(True)
        self._update_timer.timeout.connect(self._update_realtime_analysis)
        self._last_text_hash = 0

        # Set up context menu
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.customContextMenuRequested.connect(self._show_context_menu)

        # TTS setup
        self._tts_service = None
        self._tts_available = TTS_AVAILABLE
        if self._tts_available:
            self._init_tts()

    def _on_text_changed(self):
        """Handle text changes - debounce analysis for realtime modes."""
        # Use longer debounce (1.5s) to avoid freezing during rapid typing
        self._update_timer.start(1500)

    def _update_realtime_analysis(self):
        """Update analysis for any checks in realtime mode."""
        current_text = self.toPlainText()
        current_hash = hash(current_text)

        # Only update if text actually changed
        if current_hash != self._last_text_hash:
            self._last_text_hash = current_hash
            # Update the cache in overuse detector
            self.overuse_detector.update_cache(current_text)
            self.writing_highlighter.update_full_text(current_text)

            # Only rehighlight if any check is in realtime mode
            if self._has_realtime_checks():
                self.writing_highlighter.rehighlight()

    def _has_realtime_checks(self) -> bool:
        """Check if any writing check is in realtime mode."""
        return (
            self.writing_highlighter.get_spell_mode() == CheckMode.REALTIME or
            self.writing_highlighter.get_grammar_mode() == CheckMode.REALTIME or
            self.writing_highlighter.get_overuse_mode() == CheckMode.REALTIME
        )

    def do_initial_check(self):
        """Perform initial check when content is loaded."""
        current_text = self.toPlainText()
        self.overuse_detector.update_cache(current_text)
        self.writing_highlighter.update_full_text(current_text)
        self.writing_highlighter.do_initial_check()
        self._last_text_hash = hash(current_text)  # Update hash to avoid immediate re-check

    def set_spell_mode(self, mode: CheckMode):
        """Set spell checking mode."""
        self.writing_highlighter.set_spell_mode(mode)

    def set_grammar_mode(self, mode: CheckMode):
        """Set grammar checking mode."""
        self.writing_highlighter.set_grammar_mode(mode)

    def set_overuse_mode(self, mode: CheckMode):
        """Set overuse detection mode."""
        self.writing_highlighter.set_overuse_mode(mode)

    def get_spell_mode(self) -> CheckMode:
        """Get current spell checking mode."""
        return self.writing_highlighter.get_spell_mode()

    def get_grammar_mode(self) -> CheckMode:
        """Get current grammar checking mode."""
        return self.writing_highlighter.get_grammar_mode()

    def get_overuse_mode(self) -> CheckMode:
        """Get current overuse detection mode."""
        return self.writing_highlighter.get_overuse_mode()

    def _show_error_tooltip(self, global_pos: QPoint, viewport_pos: QPoint) -> bool:
        """Show tooltip for error at position. Returns True if tooltip was shown."""
        cursor = self.cursorForPosition(viewport_pos)
        block = cursor.block()
        block_number = block.blockNumber()
        position_in_block = cursor.positionInBlock()

        # Check for any error at this position
        error = self.writing_highlighter.get_error_at(block_number, position_in_block)
        if error:
            # Color-coded error type indicator
            color_map = {
                ErrorType.SPELLING: "#ff0000",  # Red
                ErrorType.GRAMMAR: "#00b400",   # Green
                ErrorType.OVERUSE: "#0078ff",   # Blue
                ErrorType.OVERUSED_PHRASE: "#9400d3",  # Purple
            }
            type_names = {
                ErrorType.SPELLING: "Spelling",
                ErrorType.GRAMMAR: "Grammar",
                ErrorType.OVERUSE: "Overused Word",
                ErrorType.OVERUSED_PHRASE: "Repeated Phrase",
            }
            color = color_map.get(error.error_type, "#000000")
            type_name = type_names.get(error.error_type, "Error")

            tooltip_text = f"<span style='color:{color}'><b>{type_name}:</b></span> {error.word}<br>{error.reason}"
            if error.suggestions:
                tooltip_text += f"<br><br><b>Suggestions:</b> {', '.join(error.suggestions[:5])}"
            tooltip_text += "<br><br><i>Right-click for options</i>"
            QToolTip.showText(global_pos, tooltip_text, self)
            return True
        else:
            QToolTip.hideText()
        return False

    def mousePressEvent(self, event: QMouseEvent):
        """Handle mouse press - Shift+click extends existing selection."""
        if (event.modifiers() & Qt.KeyboardModifier.ShiftModifier
                and event.button() == Qt.MouseButton.LeftButton
                and self.textCursor().hasSelection()):
            # Shift+click with existing selection: extend selection to click point
            click_pos = self.cursorForPosition(event.pos()).position()
            cursor = self.textCursor()
            anchor = cursor.anchor()
            cursor.setPosition(anchor)
            cursor.setPosition(click_pos, QTextCursor.MoveMode.KeepAnchor)
            self.setTextCursor(cursor)
        else:
            super().mousePressEvent(event)

    def mouseDoubleClickEvent(self, event: QMouseEvent):
        """Handle double-click for range selection.

        Double-click: select from cursor position to end of chapter.
        Shift+double-click: select from beginning of chapter to cursor position.
        """
        modifiers = event.modifiers()
        click_pos = self.cursorForPosition(event.pos()).position()

        if modifiers & Qt.KeyboardModifier.ShiftModifier:
            # Shift+double-click: select from start to click position
            cursor = self.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            cursor.setPosition(click_pos, QTextCursor.MoveMode.KeepAnchor)
            self.setTextCursor(cursor)
        else:
            # Plain double-click: select from click position to end
            cursor = self.textCursor()
            cursor.setPosition(click_pos)
            cursor.movePosition(QTextCursor.MoveOperation.End, QTextCursor.MoveMode.KeepAnchor)
            self.setTextCursor(cursor)

    def viewportEvent(self, event: QEvent) -> bool:
        """Handle viewport events including tooltip display for writing errors.

        On macOS, QTextEdit renders in a viewport, so tooltip events arrive here
        rather than in event(). This ensures tooltips work cross-platform.
        """
        if event.type() == QEvent.Type.ToolTip:
            pos = event.pos()
            global_pos = event.globalPos()
            if self._show_error_tooltip(global_pos, pos):
                return True

        return super().viewportEvent(event)

    def event(self, event: QEvent) -> bool:
        """Handle events including tooltip display for writing errors."""
        if event.type() == QEvent.Type.ToolTip:
            # Get position under cursor
            pos = event.pos()
            global_pos = event.globalPos()
            if self._show_error_tooltip(global_pos, pos):
                return True

        return super().event(event)

    def set_spell_check_enabled(self, enabled: bool):
        """Enable or disable spell checking (backward compatibility)."""
        mode = CheckMode.ON_DEMAND if enabled else CheckMode.OFF
        self.set_spell_mode(mode)
        self.spell_check_toggled.emit(enabled)

    def set_grammar_check_enabled(self, enabled: bool):
        """Enable or disable grammar checking (backward compatibility)."""
        mode = CheckMode.ON_DEMAND if enabled else CheckMode.OFF
        self.set_grammar_mode(mode)
        self.grammar_check_toggled.emit(enabled)

    def set_overuse_check_enabled(self, enabled: bool):
        """Enable or disable overuse detection (backward compatibility)."""
        mode = CheckMode.ON_DEMAND if enabled else CheckMode.OFF
        self.set_overuse_mode(mode)
        self.overuse_check_toggled.emit(enabled)

    def is_spell_check_enabled(self) -> bool:
        """Check if spell checking is enabled (not OFF)."""
        return self.get_spell_mode() != CheckMode.OFF

    def is_grammar_check_enabled(self) -> bool:
        """Check if grammar checking is enabled (not OFF)."""
        return self.get_grammar_mode() != CheckMode.OFF

    def is_overuse_check_enabled(self) -> bool:
        """Check if overuse detection is enabled (not OFF)."""
        return self.get_overuse_mode() != CheckMode.OFF

    def apply_heading(self, level: str):
        """Apply heading style to current block using Markdown prefixes.

        Args:
            level: One of "Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"
        """
        cursor = self.textCursor()

        # Select entire block
        cursor.movePosition(QTextCursor.MoveOperation.StartOfBlock)
        cursor.movePosition(QTextCursor.MoveOperation.EndOfBlock, QTextCursor.MoveMode.KeepAnchor)

        current_line = cursor.selectedText()

        # Convert level name to MarkdownStyle
        md_style = STYLE_NAMES.get(level, MarkdownStyle.NORMAL)

        # Apply the Markdown heading prefix
        new_line = apply_heading_to_line(current_line, md_style)

        # Replace the line content
        cursor.insertText(new_line)

        # Move cursor to end of the new line
        self.setTextCursor(cursor)

    def get_current_heading(self) -> str:
        """Get the heading level of the current block by detecting Markdown prefix."""
        cursor = self.textCursor()
        block = cursor.block()
        line_text = block.text()

        # Detect heading from Markdown prefix
        md_style, _ = get_line_style(line_text)
        return STYLE_TO_NAME.get(md_style, "Normal")

    def get_block_heading_level(self, block) -> HeadingLevel:
        """Get the heading level of a specific block by detecting Markdown prefix."""
        line_text = block.text()
        md_style, _ = get_line_style(line_text)

        # Map MarkdownStyle to HeadingLevel
        style_to_level = {
            MarkdownStyle.NORMAL: HeadingLevel.NORMAL,
            MarkdownStyle.TITLE: HeadingLevel.TITLE,
            MarkdownStyle.HEADING_1: HeadingLevel.HEADING_1,
            MarkdownStyle.HEADING_2: HeadingLevel.HEADING_2,
            MarkdownStyle.HEADING_3: HeadingLevel.HEADING_3,
            MarkdownStyle.HEADING_4: HeadingLevel.HEADING_4,
        }
        return style_to_level.get(md_style, HeadingLevel.NORMAL)

    def _show_context_menu(self, position):
        """Show custom context menu with error handling options."""
        menu = QMenu(self)

        # Get cursor position for error detection
        click_cursor = self.cursorForPosition(position)
        block = click_cursor.block()
        block_number = block.blockNumber()
        position_in_block = click_cursor.positionInBlock()

        # Check for any writing error at click position
        error = self.writing_highlighter.get_error_at(block_number, position_in_block)

        if error:
            # Create error-specific menu section
            type_names = {
                ErrorType.SPELLING: "Spelling",
                ErrorType.GRAMMAR: "Grammar",
                ErrorType.OVERUSE: "Overused Word",
                ErrorType.OVERUSED_PHRASE: "Repeated Phrase",
            }
            type_name = type_names.get(error.error_type, "Error")

            # Header showing error type and word
            header_label = QLabel(f"<b>{type_name}:</b> {error.word}")
            header_label.setStyleSheet("padding: 5px; background-color: #f0f0f0;")
            header_action = QWidgetAction(self)
            header_action.setDefaultWidget(header_label)
            menu.addAction(header_action)

            # Reason
            reason_label = QLabel(f"<i>{error.reason}</i>")
            reason_label.setWordWrap(True)
            reason_label.setMaximumWidth(300)
            reason_label.setStyleSheet("padding: 5px; color: #666;")
            reason_action = QWidgetAction(self)
            reason_action.setDefaultWidget(reason_label)
            menu.addAction(reason_action)

            menu.addSeparator()

            # Suggestions as clickable menu items
            if error.suggestions:
                suggestions_label = menu.addAction("Replace with:")
                suggestions_label.setEnabled(False)

                # Get cursor for word replacement
                word_cursor = self.cursorForPosition(position)
                word_cursor.select(QTextCursor.SelectionType.WordUnderCursor)

                for suggestion in error.suggestions[:7]:
                    action = menu.addAction(f"  {suggestion}")
                    action.triggered.connect(
                        lambda checked, s=suggestion, c=word_cursor: self._replace_word(c, s)
                    )

                menu.addSeparator()

            # Ignore option
            ignore_action = menu.addAction(f"Ignore '{error.word}'")
            ignore_action.triggered.connect(
                lambda: self._ignore_error(block_number, error.start, error.word, error.error_type)
            )

            # Add to dictionary (for spelling errors)
            if error.error_type == ErrorType.SPELLING:
                add_dict_action = menu.addAction(f"Add '{error.word}' to Dictionary")
                add_dict_action.triggered.connect(lambda: self._add_to_dictionary(error.word))

            menu.addSeparator()

        # Standard edit actions
        undo_action = QAction("Undo", self)
        undo_action.triggered.connect(self.undo)
        undo_action.setEnabled(self.document().isUndoAvailable())
        menu.addAction(undo_action)

        redo_action = QAction("Redo", self)
        redo_action.triggered.connect(self.redo)
        redo_action.setEnabled(self.document().isRedoAvailable())
        menu.addAction(redo_action)

        menu.addSeparator()

        cut_action = QAction("Cut", self)
        cut_action.triggered.connect(self.cut)
        cut_action.setEnabled(self.textCursor().hasSelection())
        menu.addAction(cut_action)

        copy_action = QAction("Copy", self)
        copy_action.triggered.connect(self.copy)
        copy_action.setEnabled(self.textCursor().hasSelection())
        menu.addAction(copy_action)

        paste_action = QAction("Paste", self)
        paste_action.triggered.connect(self.paste)
        menu.addAction(paste_action)

        menu.addSeparator()

        # Heading submenu
        heading_menu = menu.addMenu("Heading Style")
        for level in self.HEADING_STYLES.keys():
            action = heading_menu.addAction(level)
            action.triggered.connect(lambda checked, l=level: self.apply_heading(l))

        menu.addSeparator()

        # Context lookup menu
        cursor = self.textCursor()
        selected_text = cursor.selectedText()

        lookup_menu = menu.addMenu("Look Up Context")

        if selected_text:
            # Lookup selected text
            lookup_selected = QAction(f'Look Up "{selected_text[:30]}..."', self)
            lookup_selected.triggered.connect(lambda: self._lookup_selected_text(selected_text))
            lookup_menu.addAction(lookup_selected)
            lookup_menu.addSeparator()

        # Character lookup
        character_action = QAction("Character Reference", self)
        character_action.triggered.connect(self._lookup_character)
        lookup_menu.addAction(character_action)

        # Worldbuilding lookup
        worldbuilding_action = QAction("Worldbuilding Reference", self)
        worldbuilding_action.triggered.connect(self._lookup_worldbuilding)
        lookup_menu.addAction(worldbuilding_action)

        # Plot lookup
        plot_action = QAction("Plot Reference", self)
        plot_action.triggered.connect(self._lookup_plot)
        lookup_menu.addAction(plot_action)

        # Text-to-Speech menu
        if self.is_tts_available():
            menu.addSeparator()
            tts_menu = menu.addMenu("Read Aloud")

            # Read selection
            if selected_text:
                read_selection = QAction(f'Read Selection', self)
                read_selection.triggered.connect(self.speak_selection)
                tts_menu.addAction(read_selection)

            # Read from cursor
            read_from_cursor = QAction("Read from Cursor", self)
            read_from_cursor.triggered.connect(self.speak_from_cursor)
            tts_menu.addAction(read_from_cursor)

            # Read entire document
            read_all = QAction("Read Entire Document", self)
            read_all.triggered.connect(self.speak_document)
            tts_menu.addAction(read_all)

            tts_menu.addSeparator()

            # Stop if currently speaking
            if self.is_tts_speaking():
                stop_action = QAction("Stop Reading", self)
                stop_action.triggered.connect(self.stop_speaking)
                tts_menu.addAction(stop_action)

            # TTS Settings
            settings_action = QAction("Voice Settings...", self)
            settings_action.triggered.connect(self.show_tts_settings_dialog)
            tts_menu.addAction(settings_action)

            tts_menu.addSeparator()

            # Generate TTS Document
            if selected_text:
                gen_selection = QAction("Generate TTS Doc from Selection...", self)
                gen_selection.triggered.connect(lambda: self.show_tts_document_generator(selected_text))
                tts_menu.addAction(gen_selection)

            gen_all = QAction("Generate TTS Doc from Document...", self)
            gen_all.triggered.connect(lambda: self.show_tts_document_generator())
            tts_menu.addAction(gen_all)

        menu.exec(self.mapToGlobal(position))

    def _replace_word(self, cursor: QTextCursor, replacement: str):
        """Replace word at cursor with replacement."""
        cursor.insertText(replacement)

    def _add_to_dictionary(self, word: str):
        """Add word to custom dictionary."""
        self.spell_checker.add_word(word)
        self.writing_highlighter.rehighlight()

    def _ignore_error(self, block_number: int, start: int, word: str, error_type: ErrorType):
        """Ignore a specific error."""
        self.writing_highlighter.ignore_error(block_number, start, word, error_type)

        # For overuse errors, also ignore the word in the overuse detector
        if error_type == ErrorType.OVERUSE:
            self.overuse_detector.ignore_word(word)

    def _lookup_selected_text(self, text: str):
        """Look up context for selected text."""
        if self.lookup_context_callback:
            result = self.lookup_context_callback(text)
            dialog = ContextLookupDialog(f"Context for: {text}", result, self)
            dialog.exec()

    def _lookup_character(self):
        """Look up character reference."""
        if not self.get_character_list_callback:
            QMessageBox.information(self, "Not Available", "Character lookup not configured.")
            return

        characters = self.get_character_list_callback()
        if not characters:
            QMessageBox.information(self, "No Characters", "No characters defined in your project yet.")
            return

        dialog = QuickReferenceDialog(characters, "Character", self)
        if dialog.exec() and dialog.selected_item:
            if self.lookup_characters_callback:
                result = self.lookup_characters_callback(dialog.selected_item)
                ref_dialog = ContextLookupDialog(f"Character: {dialog.selected_item}", result, self)
                ref_dialog.exec()

    def _lookup_worldbuilding(self):
        """Look up worldbuilding reference."""
        if not self.get_worldbuilding_sections_callback:
            QMessageBox.information(self, "Not Available", "Worldbuilding lookup not configured.")
            return

        sections = self.get_worldbuilding_sections_callback()
        if not sections:
            QMessageBox.information(self, "No Worldbuilding", "No worldbuilding sections defined yet.")
            return

        dialog = QuickReferenceDialog(sections, "Worldbuilding Section", self)
        if dialog.exec() and dialog.selected_item:
            if self.lookup_worldbuilding_callback:
                result = self.lookup_worldbuilding_callback(dialog.selected_item)
                ref_dialog = ContextLookupDialog(f"Worldbuilding: {dialog.selected_item}", result, self)
                ref_dialog.exec()

    def _lookup_plot(self):
        """Look up plot reference."""
        if self.lookup_plot_callback:
            result = self.lookup_plot_callback()
            dialog = ContextLookupDialog("Plot Reference", result, self)
            dialog.exec()

    def import_from_docx(self, file_path: str) -> bool:
        """Import content from Word document with full formatting and heading recognition."""
        try:
            doc = Document(file_path)

            # Clear current content
            self.clear()

            # Import paragraphs with formatting
            cursor = self.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Start)

            for para in doc.paragraphs:
                # Check for heading styles
                style_name = para.style.name if para.style else ""

                block_format = QTextBlockFormat()
                heading_level = HeadingLevel.NORMAL

                # Handle heading styles from Word
                if style_name == "Title":
                    heading_level = HeadingLevel.TITLE
                    style_def = self.HEADING_STYLES["Title"]
                elif style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                        level_map = {
                            1: ("Heading 1", HeadingLevel.HEADING_1),
                            2: ("Heading 2", HeadingLevel.HEADING_2),
                            3: ("Heading 3", HeadingLevel.HEADING_3),
                            4: ("Heading 4", HeadingLevel.HEADING_4),
                        }
                        style_name, heading_level = level_map.get(level, ("Normal", HeadingLevel.NORMAL))
                        style_def = self.HEADING_STYLES.get(style_name, self.HEADING_STYLES["Normal"])
                    except (ValueError, IndexError):
                        style_def = self.HEADING_STYLES["Normal"]
                else:
                    style_def = self.HEADING_STYLES["Normal"]

                # Apply heading style if it's a heading
                if heading_level != HeadingLevel.NORMAL:
                    font_size, bold, top_margin, bottom_margin, _ = style_def
                    char_format = QTextCharFormat()
                    char_format.setFontPointSize(font_size)
                    char_format.setFontWeight(QFont.Weight.Bold if bold else QFont.Weight.Normal)
                    block_format.setTopMargin(top_margin)
                    block_format.setBottomMargin(bottom_margin)
                    cursor.setBlockFormat(block_format)
                    cursor.insertText(para.text, char_format)

                    # Store heading level metadata
                    cursor.block().setUserData(BlockData(heading_level))

                    cursor.insertBlock()
                    continue

                # Handle paragraph alignment
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    block_format.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    cursor.setBlockFormat(block_format)

                # Process runs (formatted text segments)
                for run in para.runs:
                    # Create text format
                    char_format = QTextCharFormat()

                    # Font family
                    if run.font.name:
                        char_format.setFontFamily(run.font.name)

                    # Font size
                    if run.font.size and run.font.size.pt and run.font.size.pt > 0:
                        char_format.setFontPointSize(run.font.size.pt)

                    # Bold
                    if run.font.bold:
                        char_format.setFontWeight(QFont.Weight.Bold)

                    # Italic
                    if run.font.italic:
                        char_format.setFontItalic(True)

                    # Underline
                    if run.font.underline:
                        char_format.setFontUnderline(True)

                    # Color
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        color = QColor(rgb[0], rgb[1], rgb[2])
                        char_format.setForeground(color)

                    # Insert text with formatting
                    cursor.insertText(run.text, char_format)

                # New paragraph
                cursor.insertBlock()

            return True

        except Exception as e:
            print(f"Error importing Word document: {e}")
            return False

    def export_to_docx(self, file_path: str, title: str = "") -> bool:
        """Export content to Word document with full formatting preservation.

        Properly exports heading styles so they appear in Word's navigation pane
        and can be used for table of contents generation.
        """
        try:
            doc = Document()

            # Add title if provided (as document title, not heading)
            if title:
                doc.add_heading(title, 0)

            # Parse document block by block
            block = self.document().begin()

            while block.isValid():
                block_text = block.text()

                if not block_text.strip():
                    # Empty paragraph
                    doc.add_paragraph("")
                    block = block.next()
                    continue

                # Get heading level and content from Markdown
                md_style, content = get_line_style(block_text)
                heading_level = self.get_block_heading_level(block)

                # Map heading levels to Word heading styles
                if heading_level == HeadingLevel.TITLE:
                    doc.add_heading(content, 0)  # Title style
                elif heading_level == HeadingLevel.HEADING_1:
                    doc.add_heading(content, 1)
                elif heading_level == HeadingLevel.HEADING_2:
                    doc.add_heading(content, 2)
                elif heading_level == HeadingLevel.HEADING_3:
                    doc.add_heading(content, 3)
                elif heading_level == HeadingLevel.HEADING_4:
                    doc.add_heading(content, 4)
                else:
                    # Regular paragraph - process Markdown inline formatting
                    para = doc.add_paragraph()
                    self._add_markdown_runs_to_paragraph(para, block_text)

                block = block.next()

            doc.save(file_path)
            return True

        except Exception as e:
            print(f"Error exporting to Word: {e}")
            return False

    def _add_markdown_runs_to_paragraph(self, para, text: str):
        """Parse Markdown inline formatting and add runs to a Word paragraph.

        Handles **bold**, *italic*, ***bold italic***, and plain text.
        """
        # Pattern to match bold-italic, bold, or italic markers
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|_(.+?)_)'
        last_end = 0

        for match in re.finditer(pattern, text):
            # Add text before this match as plain
            if match.start() > last_end:
                para.add_run(text[last_end:match.start()])

            # Determine which group matched
            if match.group(2):  # ***bold italic***
                run = para.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(3):  # **bold**
                run = para.add_run(match.group(3))
                run.bold = True
            elif match.group(4):  # *italic*
                run = para.add_run(match.group(4))
                run.italic = True
            elif match.group(5):  # _italic_
                run = para.add_run(match.group(5))
                run.italic = True

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            para.add_run(text[last_end:])

    def set_callbacks(
        self,
        lookup_worldbuilding: Optional[Callable] = None,
        lookup_characters: Optional[Callable] = None,
        lookup_plot: Optional[Callable] = None,
        lookup_context: Optional[Callable] = None,
        get_character_list: Optional[Callable] = None,
        get_worldbuilding_sections: Optional[Callable] = None
    ):
        """Set callback functions for context lookup."""
        self.lookup_worldbuilding_callback = lookup_worldbuilding
        self.lookup_characters_callback = lookup_characters
        self.lookup_plot_callback = lookup_plot
        self.lookup_context_callback = lookup_context
        self.get_character_list_callback = get_character_list
        self.get_worldbuilding_sections_callback = get_worldbuilding_sections

    # ==================== Text-to-Speech Methods ====================

    def _init_tts(self):
        """Initialize TTS service."""
        try:
            self._tts_service = get_tts_service()
            self._tts_service.set_callbacks(
                on_start=self._on_tts_start,
                on_end=self._on_tts_end,
                on_error=self._on_tts_error,
                on_progress=self._on_tts_progress
            )
            # Apply saved TTS settings
            self._apply_tts_settings()
        except Exception as e:
            print(f"Failed to initialize TTS: {e}")
            self._tts_available = False

    def _apply_tts_settings(self):
        """Apply TTS settings from config."""
        try:
            from src.config.ai_config import get_ai_config
            config = get_ai_config()
            settings = config.settings

            # Apply speech rate (default 150 WPM - normal pace)
            rate = settings.get("tts_rate", 150)
            self._tts_service.set_rate(rate)

            # Apply volume
            volume = settings.get("tts_volume", 1.0)
            self._tts_service.set_volume(volume)

            # Apply engine preference
            engine_name = settings.get("tts_engine", "system")
            try:
                engine_map = {
                    "vibevoice": TTSEngine.VIBEVOICE,
                    "edge": TTSEngine.EDGE,
                    "chatterbox": TTSEngine.CHATTERBOX,
                    "kokoro": TTSEngine.KOKORO,
                    "system": TTSEngine.SYSTEM,
                }
                engine = engine_map.get(engine_name, TTSEngine.SYSTEM)
                self._tts_service.set_engine(engine)
            except ValueError:
                pass

            # Apply voice: explicit voice setting takes priority, but only
            # if it's valid for the current engine. Otherwise fall back to
            # genre-based voice, then engine default.
            voice = settings.get("tts_voice", "")
            genre = settings.get("tts_genre", "")

            # Validate saved voice against current engine's voice list
            voice_valid = False
            if voice and voice != "default":
                try:
                    available_ids = {v.id for v in self._tts_service.get_voices()}
                    voice_valid = voice in available_ids
                except Exception:
                    voice_valid = True  # Assume valid if we can't check

            if voice == "default":
                # "System Default" — clear any explicit voice
                self._tts_service.set_voice(None)
            elif voice and voice_valid:
                self._tts_service.set_voice(voice)
            elif genre and engine_name:
                try:
                    from src.services.tts_service import get_genre_voice
                    voice_id = get_genre_voice(genre, engine_name)
                    if voice_id and voice_id != "default":
                        self._tts_service.set_voice(voice_id)
                    else:
                        self._tts_service.set_voice(None)
                except Exception:
                    self._tts_service.set_voice(None)
            else:
                # No voice or genre set — use engine default
                self._tts_service.set_voice(None)

            # Apply VibeVoice settings if configured
            vv_path = settings.get("vibevoice_path", "")
            if vv_path:
                self._tts_service.set_vibevoice_path(vv_path)

            vv_model = settings.get("vibevoice_model", "1.5B")
            self._tts_service.set_vibevoice_model(vv_model)

            vv_voice = settings.get("vibevoice_voice", "emma")
            self._tts_service.set_vibevoice_voice(vv_voice)

        except Exception as e:
            print(f"Failed to apply TTS settings: {e}")

    def _on_tts_start(self):
        """Handle TTS start event (called from background thread)."""
        # Use QTimer.singleShot to safely emit from main thread
        QTimer.singleShot(0, self.tts_started.emit)

    def _on_tts_end(self):
        """Handle TTS end event (called from background thread)."""
        # Use QTimer.singleShot to safely emit from main thread
        QTimer.singleShot(0, self.tts_stopped.emit)

    def _on_tts_error(self, error: str):
        """Handle TTS error event (called from background thread)."""
        # Use QTimer.singleShot to safely emit from main thread
        QTimer.singleShot(0, lambda: self.tts_error.emit(error))

    def _on_tts_progress(self, message: str):
        """Handle TTS progress update (for VibeVoice, called from background thread)."""
        # Use QTimer.singleShot to safely emit from main thread
        QTimer.singleShot(0, lambda: self.tts_progress.emit(message))

    def is_tts_available(self) -> bool:
        """Check if TTS is available."""
        return self._tts_available and self._tts_service is not None

    def is_tts_speaking(self) -> bool:
        """Check if TTS is currently speaking."""
        if self._tts_service:
            return self._tts_service.is_speaking
        return False

    def is_tts_paused(self) -> bool:
        """Check if TTS is currently paused."""
        if self._tts_service:
            return self._tts_service.is_paused
        return False

    def pause_speaking(self):
        """Pause TTS playback (can be resumed)."""
        if self._tts_service:
            return self._tts_service.pause()
        return False

    def resume_speaking(self):
        """Resume TTS playback after pause."""
        if self._tts_service:
            return self._tts_service.resume()
        return False

    def speak_text(self, text: str = None):
        """Speak text using TTS.

        Args:
            text: Text to speak. If None, speaks selected text or entire document.

        If playback is already running or paused, this call is refused
        (returns without starting new speech). Use stop_speaking() first
        if you need to interrupt.
        """
        if not self.is_tts_available():
            QMessageBox.warning(
                self,
                "TTS Not Available",
                "Text-to-Speech is not available.\n\n"
                "Install with: pip install pyttsx3 edge-tts"
            )
            return

        # Don't start a new playback while one is running or paused
        if self._tts_service.is_speaking:
            state = "paused" if self._tts_service.is_paused else "playing"
            QMessageBox.information(
                self, "Already Reading",
                f"Audio is currently {state}. "
                f"Stop the current playback before starting a new one.")
            return

        # Re-apply settings so engine/voice changes take effect immediately
        self._apply_tts_settings()

        if text is None:
            # Use selected text or entire document
            cursor = self.textCursor()
            if cursor.hasSelection():
                text = cursor.selectedText()
                # Replace paragraph separators with newlines
                text = text.replace('\u2029', '\n')
            else:
                text = self.toPlainText()

        if not text.strip():
            QMessageBox.information(self, "Nothing to Speak", "No text to read aloud.")
            return

        # Strip Markdown formatting before speaking
        clean_text = strip_markdown(text)

        # Optional: AI enhancement for more natural prosody (adds pauses,
        # expands abbreviations). Controlled by the tts_ai_enhance setting.
        try:
            from src.config.ai_config import get_ai_config
            settings = get_ai_config().settings
            if settings.get("tts_ai_enhance", False):
                from src.services.tts_service import enhance_text_for_speech
                genre = settings.get("tts_genre", "")
                clean_text = enhance_text_for_speech(clean_text, genre)
        except Exception:
            pass

        self._tts_service.speak(clean_text)

    def speak_selection(self):
        """Speak the currently selected text."""
        # Stop any ongoing playback first
        if self._tts_service and self._tts_service.is_speaking:
            self._tts_service.stop()

        cursor = self.textCursor()
        if cursor.hasSelection():
            text = cursor.selectedText().replace('\u2029', '\n')
            if text.strip():
                # Strip Markdown formatting before speaking
                clean_text = strip_markdown(text)
                self._tts_service.speak(clean_text)
            else:
                QMessageBox.information(self, "Nothing Selected", "Please select some text to read aloud.")
        else:
            QMessageBox.information(self, "Nothing Selected", "Please select some text to read aloud.")

    def speak_document(self):
        """Speak the entire document."""
        if self._tts_service and self._tts_service.is_speaking:
            state = "paused" if self._tts_service.is_paused else "playing"
            QMessageBox.information(
                self, "Already Reading",
                f"Audio is currently {state}. "
                f"Stop the current playback before starting a new one.")
            return
        self._apply_tts_settings()

        text = self.toPlainText()
        if text.strip():
            # Strip Markdown formatting before speaking
            clean_text = strip_markdown(text)
            self._tts_service.speak(clean_text)
        else:
            QMessageBox.information(self, "Empty Document", "The document is empty.")

    def speak_from_cursor(self):
        """Speak from cursor position to end of document."""
        if self._tts_service and self._tts_service.is_speaking:
            state = "paused" if self._tts_service.is_paused else "playing"
            QMessageBox.information(
                self, "Already Reading",
                f"Audio is currently {state}. "
                f"Stop the current playback before starting a new one.")
            return
        self._apply_tts_settings()

        cursor = self.textCursor()
        position = cursor.position()
        full_text = self.toPlainText()
        text = full_text[position:]
        if text.strip():
            # Strip Markdown formatting before speaking
            clean_text = strip_markdown(text)
            self._tts_service.speak(clean_text)
        else:
            QMessageBox.information(self, "Nothing to Speak", "No text after cursor position.")

    def stop_speaking(self):
        """Stop TTS playback."""
        if self._tts_service:
            self._tts_service.stop()

    def get_tts_voices(self) -> List:
        """Get available TTS voices."""
        if self._tts_service:
            return self._tts_service.get_voices()
        return []

    def set_tts_voice(self, voice_id: str):
        """Set the TTS voice."""
        if self._tts_service:
            self._tts_service.set_voice(voice_id)

    def set_tts_rate(self, rate: int):
        """Set TTS speech rate."""
        if self._tts_service:
            self._tts_service.set_rate(rate)

    def set_tts_volume(self, volume: float):
        """Set TTS volume (0.0 to 1.0)."""
        if self._tts_service:
            self._tts_service.set_volume(volume)

    def set_tts_engine(self, engine_name: str):
        """Set TTS engine by name."""
        if self._tts_service:
            try:
                engine_map = {
                    "system": TTSEngine.SYSTEM,
                    "edge": TTSEngine.EDGE,
                    "vibevoice": TTSEngine.VIBEVOICE,
                    "chatterbox": TTSEngine.CHATTERBOX,
                    "kokoro": TTSEngine.KOKORO,
                }
                engine = engine_map.get(engine_name, TTSEngine.SYSTEM)
                self._tts_service.set_engine(engine)
            except ValueError as e:
                QMessageBox.warning(self, "Engine Not Available", str(e))

    def save_speech_to_file(self, output_path: str, text: str = None) -> bool:
        """Save speech to an audio file.

        Args:
            output_path: Path to save the audio file
            text: Text to convert. If None, uses selected text or entire document.

        Returns:
            True if successful
        """
        if not self.is_tts_available():
            return False

        if text is None:
            cursor = self.textCursor()
            if cursor.hasSelection():
                text = cursor.selectedText().replace('\u2029', '\n')
            else:
                text = self.toPlainText()

        if not text.strip():
            return False

        return self._tts_service.speak_to_file(text, output_path)

    def show_tts_settings_dialog(self):
        """Show TTS settings dialog."""
        if not self.is_tts_available():
            QMessageBox.warning(
                self,
                "TTS Not Available",
                "Text-to-Speech is not available.\n\n"
                "Install with: pip install pyttsx3 edge-tts"
            )
            return

        dialog = TTSSettingsDialog(self._tts_service, self)
        dialog.exec()

    def show_tts_document_generator(self, text: str = None):
        """Show dialog to generate TTS document.

        Args:
            text: Text to convert. If None, uses entire document.
        """
        if not self.is_tts_available():
            QMessageBox.warning(
                self,
                "TTS Not Available",
                "Text-to-Speech is not available.\n\n"
                "Install with: pip install pyttsx3 edge-tts"
            )
            return

        if text is None:
            text = self.toPlainText()

        if not text.strip():
            QMessageBox.information(self, "No Text", "No text to generate TTS document from.")
            return

        # Strip Markdown formatting before generating TTS document
        clean_text = strip_markdown(text)

        # Get chapter name from parent if available
        chapter_name = "chapter"
        if hasattr(self, 'chapter_name'):
            chapter_name = self.chapter_name
        elif hasattr(self.parent(), 'get_current_chapter_name'):
            chapter_name = self.parent().get_current_chapter_name() or "chapter"

        dialog = TTSDocumentGeneratorDialog(clean_text, chapter_name, self)
        dialog.exec()


class TTSSettingsDialog(QDialog):
    """Dialog for configuring TTS settings."""

    def __init__(self, tts_service, parent=None):
        super().__init__(parent)
        self.tts_service = tts_service
        self.setWindowTitle("Text-to-Speech Settings")
        self.setMinimumSize(400, 350)

        layout = QVBoxLayout(self)

        # Engine selection
        engine_group = QGroupBox("TTS Engine")
        engine_layout = QVBoxLayout(engine_group)

        self.engine_combo = QComboBox()
        engines = tts_service.get_available_engines()
        for engine in engines:
            if engine.value == "system":
                label = "System Voices (Offline)"
            elif engine.value == "vibevoice":
                label = "VibeVoice (High Quality Local Neural)"
            else:
                label = "Edge Neural Voices (Online)"
            self.engine_combo.addItem(label, engine.value)
        self.engine_combo.currentIndexChanged.connect(self._on_engine_changed)
        engine_layout.addWidget(self.engine_combo)

        engine_note = QLabel(
            "System: Uses installed OS voices, works offline\n"
            "Edge: High-quality Microsoft neural voices, requires internet\n"
            "VibeVoice: Best quality local TTS (requires installation)"
        )
        engine_note.setStyleSheet("color: #6b7280; font-size: 10px;")
        engine_note.setWordWrap(True)
        engine_layout.addWidget(engine_note)

        layout.addWidget(engine_group)

        # Voice selection
        voice_group = QGroupBox("Voice")
        voice_layout = QVBoxLayout(voice_group)

        self.voice_combo = QComboBox()
        self._populate_voices()
        self.voice_combo.currentIndexChanged.connect(self._on_voice_changed)
        voice_layout.addWidget(self.voice_combo)

        layout.addWidget(voice_group)

        # Rate control
        rate_group = QGroupBox("Speech Rate")
        rate_layout = QHBoxLayout(rate_group)

        self.rate_slider = QSlider(Qt.Orientation.Horizontal)
        self.rate_slider.setRange(50, 300)
        self.rate_slider.setValue(150)
        self.rate_slider.valueChanged.connect(self._on_rate_changed)
        rate_layout.addWidget(self.rate_slider)

        self.rate_label = QLabel("150 wpm")
        self.rate_label.setMinimumWidth(60)
        rate_layout.addWidget(self.rate_label)

        layout.addWidget(rate_group)

        # Volume control
        volume_group = QGroupBox("Volume")
        volume_layout = QHBoxLayout(volume_group)

        self.volume_slider = QSlider(Qt.Orientation.Horizontal)
        self.volume_slider.setRange(0, 100)
        self.volume_slider.setValue(100)
        self.volume_slider.valueChanged.connect(self._on_volume_changed)
        volume_layout.addWidget(self.volume_slider)

        self.volume_label = QLabel("100%")
        self.volume_label.setMinimumWidth(40)
        volume_layout.addWidget(self.volume_label)

        layout.addWidget(volume_group)

        # Test button
        test_layout = QHBoxLayout()
        test_btn = QPushButton("Test Voice")
        test_btn.clicked.connect(self._test_voice)
        test_layout.addWidget(test_btn)

        stop_btn = QPushButton("Stop")
        stop_btn.clicked.connect(self.tts_service.stop)
        test_layout.addWidget(stop_btn)

        test_layout.addStretch()
        layout.addLayout(test_layout)

        layout.addStretch()

        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

    def _populate_voices(self):
        """Populate voice combo box."""
        self.voice_combo.clear()
        voices = self.tts_service.get_voices()
        for voice in voices:
            self.voice_combo.addItem(f"{voice.name} ({voice.language})", voice.id)

    def _on_engine_changed(self, index):
        """Handle engine selection change."""
        engine_value = self.engine_combo.currentData()
        try:
            engine_map = {
                "system": TTSEngine.SYSTEM,
                "edge": TTSEngine.EDGE,
                "vibevoice": TTSEngine.VIBEVOICE,
                "chatterbox": TTSEngine.CHATTERBOX,
            }
            engine = engine_map.get(engine_value, TTSEngine.EDGE)
            self.tts_service.set_engine(engine)
            self._populate_voices()
        except ValueError as e:
            QMessageBox.warning(self, "Engine Error", str(e))

    def _on_voice_changed(self, index):
        """Handle voice selection change."""
        voice_id = self.voice_combo.currentData()
        if voice_id:
            self.tts_service.set_voice(voice_id)

    def _on_rate_changed(self, value):
        """Handle rate slider change."""
        self.rate_label.setText(f"{value} wpm")
        self.tts_service.set_rate(value)

    def _on_volume_changed(self, value):
        """Handle volume slider change."""
        self.volume_label.setText(f"{value}%")
        self.tts_service.set_volume(value / 100.0)

    def _test_voice(self):
        """Test the current voice."""
        self.tts_service.speak(
            "Hello! This is a test of the text to speech system. "
            "I can read your writing aloud so you can hear how it sounds."
        )


class TTSDocumentGeneratorDialog(QDialog):
    """Dialog for generating TTS documents with speaker configuration."""

    # Signal emitted when TTS document is generated
    document_generated = pyqtSignal(str)  # Path to generated file

    def __init__(self, text: str, chapter_name: str = "chapter", parent=None):
        super().__init__(parent)
        self.text = text
        self.chapter_name = chapter_name
        self.speaker_configs: List[SpeakerConfig] = []

        self.setWindowTitle("Generate TTS Document")
        self.setMinimumSize(550, 500)

        self._init_ui()

    def _init_ui(self):
        layout = QVBoxLayout(self)

        # Info header
        info_label = QLabel(
            "Generate a TTS-formatted document for text-to-speech synthesis.\n"
            "Configure speakers/voices for narrator and dialogue."
        )
        info_label.setWordWrap(True)
        info_label.setStyleSheet("color: #374151; padding: 8px; background-color: #f3f4f6; border-radius: 4px;")
        layout.addWidget(info_label)

        # Format selection
        format_group = QGroupBox("Output Format")
        format_layout = QVBoxLayout(format_group)

        self.format_combo = QComboBox()
        self.format_combo.addItem("VibeVoice (Speaker N: text)", "vibevoice")
        self.format_combo.addItem("Plain Text", "plain")
        self.format_combo.addItem("SSML (Speech Synthesis Markup)", "ssml")
        format_layout.addWidget(self.format_combo)

        layout.addWidget(format_group)

        # Speaker Configuration
        speakers_group = QGroupBox("Speaker & Voice Configuration")
        speakers_layout = QVBoxLayout(speakers_group)

        # Info about voices
        voice_info = QLabel(
            "VibeVoice supports 7 distinct voices: Carter, Davis, Emma, Frank, Grace, Mike, Samuel.\n"
            "Double-click a speaker to change their name or voice."
        )
        voice_info.setWordWrap(True)
        voice_info.setStyleSheet("color: #6b7280; font-size: 11px; padding: 4px;")
        speakers_layout.addWidget(voice_info)

        # Number of speakers
        num_speakers_layout = QHBoxLayout()
        num_speakers_layout.addWidget(QLabel("Number of Speakers:"))
        self.num_speakers_combo = QComboBox()
        self.num_speakers_combo.addItems(["1 (Narrator only)", "2 (Narrator + 1 voice)", "3", "4 (Maximum)"])
        self.num_speakers_combo.setCurrentIndex(1)  # Default to 2
        self.num_speakers_combo.currentIndexChanged.connect(self._on_num_speakers_changed)
        num_speakers_layout.addWidget(self.num_speakers_combo)
        num_speakers_layout.addStretch()
        speakers_layout.addLayout(num_speakers_layout)

        # Speaker list with better formatting
        self.speaker_list = QListWidget()
        self.speaker_list.setMaximumHeight(160)
        self.speaker_list.setStyleSheet("""
            QListWidget::item {
                padding: 6px;
                border-bottom: 1px solid #e5e7eb;
            }
            QListWidget::item:selected {
                background-color: #e0e7ff;
                color: #3730a3;
            }
        """)
        self.speaker_list.itemDoubleClicked.connect(self._edit_speaker)
        speakers_layout.addWidget(self.speaker_list)

        # Edit speaker button
        edit_btn_layout = QHBoxLayout()
        edit_speaker_btn = QPushButton("Edit Selected Speaker...")
        edit_speaker_btn.clicked.connect(self._edit_speaker)
        edit_btn_layout.addWidget(edit_speaker_btn)
        edit_btn_layout.addStretch()
        speakers_layout.addLayout(edit_btn_layout)

        layout.addWidget(speakers_group)

        # Options
        options_group = QGroupBox("Options")
        options_layout = QVBoxLayout(options_group)

        from PyQt6.QtWidgets import QCheckBox
        self.dialogue_detection_check = QCheckBox("Auto-detect dialogue (quotes)")
        self.dialogue_detection_check.setChecked(True)
        options_layout.addWidget(self.dialogue_detection_check)

        self.preserve_paragraphs_check = QCheckBox("Preserve paragraph breaks")
        self.preserve_paragraphs_check.setChecked(True)
        options_layout.addWidget(self.preserve_paragraphs_check)

        layout.addWidget(options_group)

        # Output info
        output_group = QGroupBox("Output")
        output_layout = QVBoxLayout(output_group)

        output_dir = get_tts_output_dir() if TTS_AVAILABLE else Path.home() / ".writer_platform" / "tts_output"
        self.output_label = QLabel(f"Output folder: {output_dir}")
        self.output_label.setWordWrap(True)
        self.output_label.setStyleSheet("color: #6b7280; font-size: 11px;")
        output_layout.addWidget(self.output_label)

        safe_name = re.sub(r'[^\w\-_]', '_', self.chapter_name)
        self.filename_label = QLabel(f"Filename: {safe_name}_tts.txt")
        self.filename_label.setStyleSheet("color: #059669; font-weight: bold;")
        output_layout.addWidget(self.filename_label)

        layout.addWidget(output_group)

        layout.addStretch()

        # Buttons
        button_layout = QHBoxLayout()

        generate_btn = QPushButton("Generate TTS Document")
        generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #6366f1;
                color: white;
                padding: 10px 20px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #4f46e5;
            }
        """)
        generate_btn.clicked.connect(self._generate_document)
        button_layout.addWidget(generate_btn)

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)

        layout.addLayout(button_layout)

        # Initialize speakers
        self._on_num_speakers_changed(1)

    def _on_num_speakers_changed(self, index):
        """Update speaker list when number changes."""
        num_speakers = index + 1

        # Voice display names mapping
        voice_display = {
            "carter": "Carter (Male)",
            "davis": "Davis (Male)",
            "emma": "Emma (Female)",
            "frank": "Frank (Male)",
            "grace": "Grace (Female)",
            "mike": "Mike (Male)",
            "samuel": "Samuel (Male)",
        }

        # Default voice mappings - varied for different speakers
        default_voices = [
            ("Narrator", "emma"),        # Female narrator
            ("Character 1", "carter"),   # Male character
            ("Character 2", "grace"),    # Female character
            ("Character 3", "frank"),    # Male character
        ]

        self.speaker_configs = []
        self.speaker_list.clear()

        for i in range(num_speakers):
            name, voice = default_voices[i] if i < len(default_voices) else (f"Speaker {i+1}", "emma")
            config = SpeakerConfig(
                speaker_id=i + 1,
                name=name,
                voice_id=voice,
                description=""
            )
            self.speaker_configs.append(config)
            voice_name = voice_display.get(voice, voice)
            self.speaker_list.addItem(f"Speaker {i+1}: {name} - {voice_name}")

    def _edit_speaker(self, item=None):
        """Edit the selected speaker configuration."""
        row = self.speaker_list.currentRow()
        if row < 0:
            QMessageBox.information(self, "Select Speaker", "Please select a speaker to edit.")
            return

        config = self.speaker_configs[row]

        # Voice display names mapping
        voice_display = {
            "carter": "Carter (Male)",
            "davis": "Davis (Male)",
            "emma": "Emma (Female)",
            "frank": "Frank (Male)",
            "grace": "Grace (Female)",
            "mike": "Mike (Male)",
            "samuel": "Samuel (Male)",
        }

        # Show edit dialog
        dialog = SpeakerEditDialog(config, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            # Update config
            config.name = dialog.name_edit.text()
            config.voice_id = dialog.voice_combo.currentData()
            self.speaker_configs[row] = config
            voice_name = voice_display.get(config.voice_id, config.voice_id)
            self.speaker_list.item(row).setText(
                f"Speaker {config.speaker_id}: {config.name} - {voice_name}"
            )

    def _generate_document(self):
        """Generate the TTS document."""
        if not TTS_AVAILABLE:
            QMessageBox.warning(self, "TTS Not Available", "TTS services are not available.")
            return

        # Build configuration
        format_value = self.format_combo.currentData()
        tts_format = TTSFormat.VIBEVOICE
        if format_value == "plain":
            tts_format = TTSFormat.PLAIN
        elif format_value == "ssml":
            tts_format = TTSFormat.SSML

        config = TTSDocumentConfig(
            format=tts_format,
            speakers=self.speaker_configs,
            default_speaker=1,
            narrator_speaker=1,
            dialogue_detection=self.dialogue_detection_check.isChecked(),
            preserve_paragraphs=self.preserve_paragraphs_check.isChecked()
        )

        # Generate document
        generator = TTSDocumentGenerator(config)
        output_dir = get_tts_output_dir()

        try:
            filepath, speakers_used = generator.generate_tts_document(
                self.text,
                output_dir,
                self.chapter_name
            )

            # Get speaker names for VibeVoice command
            speaker_names = " ".join(s.name for s in self.speaker_configs)

            QMessageBox.information(
                self,
                "TTS Document Generated",
                f"Document saved to:\n{filepath}\n\n"
                f"Speakers used: {', '.join(speakers_used)}\n\n"
                f"To use with VibeVoice:\n"
                f"python demo/inference_from_file.py \\\n"
                f"  --model_path vibevoice/VibeVoice-1.5B \\\n"
                f"  --txt_path \"{filepath}\" \\\n"
                f"  --speaker_names {speaker_names}"
            )

            self.document_generated.emit(str(filepath))
            self.accept()

        except Exception as e:
            QMessageBox.critical(self, "Generation Failed", f"Failed to generate TTS document:\n{str(e)}")


class SpeakerEditDialog(QDialog):
    """Dialog for editing a speaker configuration."""

    # VibeVoice voice options with descriptions
    VIBEVOICE_VOICES = [
        ("carter", "Carter", "Male", "Deep, authoritative male voice - good for narrators"),
        ("davis", "Davis", "Male", "Warm, friendly male voice - good for dialogue"),
        ("emma", "Emma", "Female", "Clear, professional female voice - versatile"),
        ("frank", "Frank", "Male", "Mature, steady male voice - good for older characters"),
        ("grace", "Grace", "Female", "Soft, gentle female voice - good for younger characters"),
        ("mike", "Mike", "Male", "Energetic, youthful male voice - good for action"),
        ("samuel", "Samuel", "Male", "Distinguished, formal male voice - good for authority figures"),
    ]

    def __init__(self, config: 'SpeakerConfig', parent=None):
        super().__init__(parent)
        self.config = config

        self.setWindowTitle(f"Edit Speaker {config.speaker_id}")
        self.setMinimumWidth(450)

        layout = QVBoxLayout(self)

        # Header
        header = QLabel(f"Configure Speaker {config.speaker_id}")
        header.setStyleSheet("font-size: 14px; font-weight: bold; padding: 8px;")
        layout.addWidget(header)

        # Name
        from PyQt6.QtWidgets import QFormLayout, QLineEdit
        form_layout = QFormLayout()

        self.name_edit = QLineEdit()
        self.name_edit.setText(config.name)
        self.name_edit.setPlaceholderText("e.g., Narrator, Alice, Bob")
        form_layout.addRow("Speaker Name:", self.name_edit)

        layout.addLayout(form_layout)

        # Voice selection group
        voice_group = QGroupBox("VibeVoice Voice Selection")
        voice_layout = QVBoxLayout(voice_group)

        # Voice combo
        self.voice_combo = QComboBox()
        for voice_id, name, gender, desc in self.VIBEVOICE_VOICES:
            self.voice_combo.addItem(f"{name} ({gender})", voice_id)

        # Set current
        for i in range(self.voice_combo.count()):
            if self.voice_combo.itemData(i) == config.voice_id:
                self.voice_combo.setCurrentIndex(i)
                break

        self.voice_combo.currentIndexChanged.connect(self._on_voice_changed)
        voice_layout.addWidget(self.voice_combo)

        # Voice description
        self.voice_desc_label = QLabel()
        self.voice_desc_label.setWordWrap(True)
        self.voice_desc_label.setStyleSheet(
            "color: #6b7280; font-size: 11px; padding: 8px; "
            "background-color: #f3f4f6; border-radius: 4px;"
        )
        voice_layout.addWidget(self.voice_desc_label)
        self._on_voice_changed(self.voice_combo.currentIndex())

        # Preview button
        preview_layout = QHBoxLayout()
        preview_btn = QPushButton("Preview Voice")
        preview_btn.clicked.connect(self._preview_voice)
        preview_layout.addWidget(preview_btn)

        stop_btn = QPushButton("Stop")
        stop_btn.clicked.connect(self._stop_preview)
        preview_layout.addWidget(stop_btn)

        preview_layout.addStretch()
        voice_layout.addLayout(preview_layout)

        layout.addWidget(voice_group)

        # Voice info
        info_label = QLabel(
            "Tip: Choose different voices for narrator vs characters to make dialogue clearer."
        )
        info_label.setWordWrap(True)
        info_label.setStyleSheet("color: #059669; font-size: 11px; padding: 8px;")
        layout.addWidget(info_label)

        layout.addStretch()

        # Buttons
        button_layout = QHBoxLayout()
        ok_btn = QPushButton("OK")
        ok_btn.setDefault(True)
        ok_btn.clicked.connect(self.accept)
        button_layout.addWidget(ok_btn)

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)

        layout.addLayout(button_layout)

    def _on_voice_changed(self, index):
        """Update voice description when selection changes."""
        if index >= 0 and index < len(self.VIBEVOICE_VOICES):
            voice_id, name, gender, desc = self.VIBEVOICE_VOICES[index]
            self.voice_desc_label.setText(f"{desc}")

    def _preview_voice(self):
        """Preview the selected voice."""
        if not TTS_AVAILABLE:
            QMessageBox.warning(self, "TTS Not Available", "TTS preview is not available.")
            return

        try:
            tts_service = get_tts_service()
            voice_id = self.voice_combo.currentData()
            speaker_name = self.name_edit.text() or "Speaker"

            # Save current settings
            original_engine = tts_service._current_engine
            original_voice = tts_service._vibevoice_voice

            # Set to VibeVoice with selected voice
            if tts_service.is_vibevoice_available():
                tts_service.set_engine(TTSEngine.VIBEVOICE)
                tts_service.set_vibevoice_voice(voice_id)
                tts_service.speak(
                    f"Hello, I am {speaker_name}. This is how I sound when reading your story."
                )
            else:
                # Fall back to edge-tts for preview
                QMessageBox.information(
                    self,
                    "VibeVoice Not Installed",
                    f"VibeVoice is not installed. Voice '{voice_id}' will be used when generating the TTS document.\n\n"
                    "Install VibeVoice from Settings > Text-to-Speech to enable voice preview."
                )
        except Exception as e:
            QMessageBox.warning(self, "Preview Error", f"Could not preview voice: {e}")

    def _stop_preview(self):
        """Stop voice preview."""
        if TTS_AVAILABLE:
            try:
                tts_service = get_tts_service()
                tts_service.stop()
            except:
                pass
